"""
EduPilot AI Faculty Assistant
Module: rag.py
Version: 4.1.0
Author: EduPilot Team
Purpose: Retrieval-Augmented Generation (RAG) layer. Handles ingestion of
         knowledge-base documents (PDF, DOCX, TXT) into a FAISS vector store
         using HuggingFace sentence-transformers embeddings, and provides
         semantic retrieval with source attribution.

         Key design decisions:
         - Embedding model: all-MiniLM-L6-v2 (sentence-transformers)
         - Chunking: RecursiveCharacterTextSplitter, chunk_size=500, overlap=50
         - DOCX support included via python-docx
         - Lazy loading: the FAISS index is loaded/built on first retrieve() call
         - Index version-stamped in vectorstore/meta.json; mismatch triggers
           a warning and optional rebuild
         - Empty/missing knowledge dir is handled gracefully (returns 0 / [])
         - force_rebuild writes to a temp location then swaps to avoid corruption
"""

from __future__ import annotations

import gc
import json
import re
import shutil
import sys
import tempfile
from pathlib import Path
from typing import List, Optional, Tuple

from config import config
from logging_utils import get_logger
from models import Question, SourceAttribution

logger = get_logger(__name__)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS: tuple[str, ...] = (".pdf", ".txt", ".docx")

# ---------------------------------------------------------------------------
# Process-wide embeddings cache — avoids reloading the ONNX model on every
# RAGModule instantiation (one model load per process, not per pipeline run).
# ---------------------------------------------------------------------------
_EMBEDDINGS_CACHE: dict = {}

# Number of chunks embedded per FAISS.from_documents()/add_documents() call
# during ingest(). Keeps peak memory roughly constant regardless of
# knowledge-base size — see the comment in ingest() for why this matters
# on memory-constrained hosts.
EMBEDDING_BATCH_SIZE = 32
CHUNK_SIZE: int = 500
CHUNK_OVERLAP: int = 50
META_FILENAME: str = "meta.json"
INDEX_FILENAME: str = "index.faiss"
PKL_FILENAME: str = "index.pkl"


# ---------------------------------------------------------------------------
# Custom exceptions
# ---------------------------------------------------------------------------

class KnowledgeBaseEmptyError(RuntimeError):
    """Raised when retrieval is attempted on an empty or missing index."""


class RAGIngestionError(RuntimeError):
    """Raised when document ingestion fails unrecoverably."""


# ---------------------------------------------------------------------------
# Document loading helpers
# ---------------------------------------------------------------------------

def _load_pdf(path: Path) -> List[Tuple[str, int]]:
    """Load a PDF file and return (text, page_number) pairs.

    Args:
        path: Absolute path to the PDF file.

    Returns:
        List of (page_text, 1-based page number) tuples. Empty pages are
        skipped.

    Raises:
        RAGIngestionError: If the file cannot be parsed.
    """
    try:
        from pypdf import PdfReader  # lazy import — keep startup fast
        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                pages.append((text, i))
        logger.debug("Loaded PDF %s: %d non-empty pages", path.name, len(pages))
        return pages
    except Exception as exc:
        raise RAGIngestionError(f"Failed to load PDF '{path}': {exc}") from exc


def _load_txt(path: Path) -> List[Tuple[str, int]]:
    """Load a plain-text file and return it as a single (text, 0) pair.

    Args:
        path: Absolute path to the text file.

    Returns:
        List containing a single (full_text, 0) tuple (page 0 = n/a).

    Raises:
        RAGIngestionError: If the file cannot be read.
    """
    try:
        text = path.read_text(encoding="utf-8", errors="replace").strip()
        if not text:
            return []
        logger.debug("Loaded TXT %s: %d chars", path.name, len(text))
        return [(text, 0)]
    except Exception as exc:
        raise RAGIngestionError(f"Failed to load TXT '{path}': {exc}") from exc


def _load_docx(path: Path) -> List[Tuple[str, int]]:
    """Load a DOCX file and return paragraph text as a single (text, 0) pair.

    Args:
        path: Absolute path to the DOCX file.

    Returns:
        List containing a single (full_text, 0) tuple.

    Raises:
        RAGIngestionError: If the file cannot be parsed.
    """
    try:
        from docx import Document as DocxDocument  # lazy import
        doc = DocxDocument(str(path))
        text = "\n".join(
            para.text for para in doc.paragraphs if para.text.strip()
        ).strip()
        if not text:
            return []
        logger.debug("Loaded DOCX %s: %d chars", path.name, len(text))
        return [(text, 0)]
    except Exception as exc:
        raise RAGIngestionError(f"Failed to load DOCX '{path}': {exc}") from exc


def _load_document(path: Path) -> List[Tuple[str, int]]:
    """Dispatch to the appropriate loader based on file extension.

    Args:
        path: Path to the document.

    Returns:
        List of (text, page_number) tuples.

    Raises:
        RAGIngestionError: For unsupported or unreadable files.
    """
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _load_pdf(path)
    elif ext == ".txt":
        return _load_txt(path)
    elif ext == ".docx":
        return _load_docx(path)
    else:
        raise RAGIngestionError(f"Unsupported file type: '{path.suffix}'")


# ---------------------------------------------------------------------------
# Chunking helpers
# ---------------------------------------------------------------------------

def _chunk_text(
    text: str,
    source_name: str,
    page_number: int,
    chunk_size: int = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> List[dict]:
    """Split text into overlapping chunks with metadata.

    Args:
        text: Raw text to split.
        source_name: Human-readable document name (filename) for attribution.
        page_number: Originating page (0 if N/A for TXT/DOCX).
        chunk_size: Maximum characters per chunk.
        chunk_overlap: Overlap in characters between successive chunks.

    Returns:
        List of dicts with keys: ``text``, ``source``, ``page``,
        ``chunk_index``.
    """
    from langchain_text_splitters import RecursiveCharacterTextSplitter  # lazy

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
    )
    raw_chunks = splitter.split_text(text)
    return [
        {
            "text": chunk,
            "source": source_name,
            "page": page_number,
            "chunk_index": idx,
        }
        for idx, chunk in enumerate(raw_chunks)
    ]


# ---------------------------------------------------------------------------
# Format helper (module-level, used by downstream agents)
# ---------------------------------------------------------------------------

def format_rag_context(
    results: List[Tuple[str, SourceAttribution]],
    max_chunks: int = 5,
) -> str:
    """Serialise retrieve() results into a prompt-ready context block.

    Each chunk is presented with its source reference so the LLM can
    attribute generated questions to specific documents.

    Args:
        results: Output of :meth:`RAGModule.retrieve` — a list of
            ``(chunk_text, SourceAttribution)`` tuples ordered by
            relevance descending.
        max_chunks: Maximum number of chunks to include (safety cap).

    Returns:
        A formatted multi-line string ready for injection into an LLM prompt.
        Returns an empty string when *results* is empty.

    Example output::

        [Source 1: lecture_notes.txt, page 0, score 0.87]
        Arrays store elements of the same type...

        [Source 2: syllabus.pdf, page 3, score 0.82]
        The course covers data structures including...
    """
    if not results:
        return ""

    parts: List[str] = []
    for i, (chunk_text, attr) in enumerate(results[:max_chunks], start=1):
        page_label = f", page {attr.page_number}" if attr.page_number else ""
        header = (
            f"[Source {i}: {attr.document_name}{page_label}, "
            f"score {attr.relevance_score:.2f}]"
        )
        parts.append(f"{header}\n{chunk_text.strip()}")

    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# RAGModule
# ---------------------------------------------------------------------------

class RAGModule:
    """Manages the FAISS vector store and retrieval pipeline.

    The RAG module indexes documents from the knowledge directory and
    retrieves relevant chunks at query time. Each retrieved chunk carries
    source attribution metadata used downstream for display and export.

    Usage::

        rag = RAGModule()
        # Ingest all documents in knowledge/ (no-op if already indexed):
        count = rag.ingest()
        # Retrieve top-5 chunks for a query:
        results = rag.retrieve("binary search tree traversal", top_k=5)
        context_str = format_rag_context(results)
    """

    def __init__(self) -> None:
        """Initialise the RAG module.

        Reads vectorstore and knowledge paths from the central config.
        Does not load the vector store eagerly; the index is loaded/built
        lazily on the first :meth:`retrieve` call.
        """
        self._vectorstore_path: Path = config.vectorstore_path
        self._knowledge_dir: Path = config.knowledge_dir
        self._embedding_model_name: str = config.embedding_model_name
        self._store = None  # populated lazily by load_or_create()
        logger.debug(
            "RAGModule initialised. vectorstore=%s knowledge=%s embedding=%s",
            self._vectorstore_path,
            self._knowledge_dir,
            self._embedding_model_name,
        )

    # ------------------------------------------------------------------
    # Meta / version stamp helpers
    # ------------------------------------------------------------------

    def _meta_path(self) -> Path:
        """Return the path to the vectorstore metadata file.

        Returns:
            Path: ``<vectorstore_path>/meta.json``
        """
        return self._vectorstore_path / META_FILENAME

    def _read_meta(self) -> dict:
        """Read and return the vectorstore metadata dict.

        Returns:
            dict with at least ``{"embedding_model": str}``, or empty dict
            if the file does not exist.
        """
        meta_path = self._meta_path()
        if meta_path.exists():
            try:
                return json.loads(meta_path.read_text(encoding="utf-8"))
            except Exception as exc:
                logger.warning("Could not read vectorstore meta.json: %s", exc)
        return {}

    def _write_meta(self) -> None:
        """Persist the current embedding model name to meta.json.

        Writes ``{"embedding_model": "<model_name>"}`` to the vectorstore
        directory for version-stamp checking on subsequent loads.
        """
        self._vectorstore_path.mkdir(parents=True, exist_ok=True)
        meta = {"embedding_model": self._embedding_model_name}
        self._meta_path().write_text(
            json.dumps(meta, indent=2), encoding="utf-8"
        )
        logger.debug("Wrote vectorstore meta.json: %s", meta)

    def _check_meta(self) -> bool:
        """Check whether the persisted index was built with the current model.

        Returns:
            ``True`` if the embedding model matches (or no meta exists yet).
            ``False`` on a mismatch — caller should warn and consider rebuild.
        """
        meta = self._read_meta()
        stored_model = meta.get("embedding_model")
        if stored_model and stored_model != self._embedding_model_name:
            logger.warning(
                "Embedding model mismatch: index was built with '%s' but "
                "current config is '%s'. Retrieval quality may be degraded. "
                "Run ingest(force_rebuild=True) to rebuild.",
                stored_model,
                self._embedding_model_name,
            )
            return False
        return True

    # ------------------------------------------------------------------
    # Index persistence helpers
    # ------------------------------------------------------------------

    def _clear_index(self) -> None:
        """Remove any persisted index files and drop the in-memory store.

        Called on forced rebuilds that end with zero indexable content so a
        stale index can never keep serving deleted documents.
        """
        self._store = None
        for fname in (INDEX_FILENAME, PKL_FILENAME, META_FILENAME):
            path = self._vectorstore_path / fname
            try:
                if path.exists():
                    path.unlink()
                    logger.info("Removed stale vectorstore file: %s", path)
            except OSError as exc:
                logger.warning("Could not remove %s: %s", path, exc)

    def _index_exists(self) -> bool:
        """Return True if a valid FAISS index is present on disk.

        Returns:
            bool: True when both ``index.faiss`` and ``index.pkl`` exist.
        """
        return (
            (self._vectorstore_path / INDEX_FILENAME).exists()
            and (self._vectorstore_path / PKL_FILENAME).exists()
        )

    def _get_embeddings(self):
        """Construct and return the FastEmbed embeddings object.

        Uses ONNX-based FastEmbed (no torch dependency — torch cannot be
        installed in this environment). Downloads the model on first call
        (~65 MB for bge-small-en-v1.5); subsequent calls use the local cache.

        The returned instance is cached in ``_EMBEDDINGS_CACHE`` keyed by
        model name so the ONNX runtime is loaded only once per process,
        regardless of how many RAGModule instances are created.

        Returns:
            FastEmbedEmbeddings instance.
        """
        if self._embedding_model_name in _EMBEDDINGS_CACHE:
            return _EMBEDDINGS_CACHE[self._embedding_model_name]

        from langchain_community.embeddings.fastembed import (  # lazy import
            FastEmbedEmbeddings,
        )
        logger.info(
            "Loading embedding model '%s' (may download on first run)…",
            self._embedding_model_name,
        )
        # Use an explicit, absolute cache directory (not the library default,
        # which is a *relative* "local_cache" resolved against whatever the
        # current working directory happens to be at call time). This
        # guarantees the build step and the running app process always
        # agree on where the cached model lives, regardless of CWD
        # differences between build and runtime on the host.
        cache_dir = str(Path(__file__).resolve().parent / "model_cache")
        emb = FastEmbedEmbeddings(
            model_name=self._embedding_model_name, cache_dir=cache_dir
        )
        _EMBEDDINGS_CACHE[self._embedding_model_name] = emb
        return emb

    # ------------------------------------------------------------------
    # Public interface
    # ------------------------------------------------------------------

    def load_or_create(self, force_rebuild: bool = False) -> None:
        """Load the FAISS index from disk, or ingest documents to create it.

        If ``force_rebuild`` is True, the existing index is discarded and
        re-built from the knowledge directory regardless.

        Args:
            force_rebuild: When ``True``, rebuild the index unconditionally.
        """
        if force_rebuild:
            logger.info("force_rebuild=True — discarding existing index.")
            self.ingest(force_rebuild=True)
            return

        if self._store is not None:
            logger.debug("Index already loaded in memory, skipping reload.")
            return

        if self._index_exists():
            if not self._check_meta():
                # Index was built with a different embedding model — loading
                # it would give semantically wrong similarity scores.
                logger.info(
                    "Embedding model changed — rebuilding index automatically."
                )
                self.ingest(force_rebuild=True)
                return
            try:
                from langchain_community.vectorstores import FAISS  # lazy
                embeddings = self._get_embeddings()
                self._store = FAISS.load_local(
                    str(self._vectorstore_path),
                    embeddings,
                    allow_dangerous_deserialization=True,
                )
                logger.info(
                    "FAISS index loaded from disk: %s",
                    self._vectorstore_path,
                )
                return
            except Exception as exc:
                logger.warning(
                    "Failed to load existing index (%s); rebuilding.", exc
                )
                self._store = None

        # No index on disk (or load failed) — ingest from scratch.
        self.ingest(force_rebuild=False)

    def ingest(self, force_rebuild: bool = False) -> int:
        """Ingest all documents from the knowledge directory into FAISS.

        Scans ``knowledge/`` for PDF, DOCX, and TXT files, chunks them,
        embeds them with the configured sentence-transformer model, and
        writes the FAISS index to ``vectorstore/``.

        When ``force_rebuild=True`` the new index is written to a temporary
        directory first, then atomically swapped into place so the existing
        index is never left in a corrupt state mid-write.

        Args:
            force_rebuild: When ``True``, rebuild the index even if it
                already exists on disk.

        Returns:
            int: Number of document chunks indexed. Returns 0 gracefully
            when the knowledge directory is missing or empty.
        """
        # ── Guard: skip if index already exists and no rebuild requested ──
        if not force_rebuild and self._index_exists() and self._store is not None:
            logger.info("Index already in memory and on disk; skipping ingest.")
            return 0

        # ── Scan knowledge directory ──────────────────────────────────────
        if not self._knowledge_dir.exists():
            logger.warning(
                "Knowledge directory does not exist: %s. "
                "Create it and add documents, then call ingest() again.",
                self._knowledge_dir,
            )
            if force_rebuild:
                self._clear_index()
            return 0

        doc_paths = sorted(
            p
            for p in self._knowledge_dir.iterdir()
            if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
        )

        if not doc_paths:
            logger.warning(
                "Knowledge directory '%s' contains no supported documents "
                "(%s). Skipping index build.",
                self._knowledge_dir,
                ", ".join(SUPPORTED_EXTENSIONS),
            )
            if force_rebuild:
                # A forced rebuild with zero documents must not leave a stale
                # index behind — retrieval would silently serve old content.
                self._clear_index()
            return 0

        logger.info(
            "Ingesting %d document(s) from %s …",
            len(doc_paths),
            self._knowledge_dir,
        )

        # ── Load and chunk all documents ─────────────────────────────────
        all_chunks: List[dict] = []
        for path in doc_paths:
            try:
                pages = _load_document(path)
            except RAGIngestionError as exc:
                logger.error("Skipping '%s': %s", path.name, exc)
                continue

            for text, page_num in pages:
                chunks = _chunk_text(
                    text,
                    source_name=path.name,
                    page_number=page_num,
                )
                all_chunks.extend(chunks)
                logger.debug(
                    "  %s page %d → %d chunks", path.name, page_num, len(chunks)
                )

        if not all_chunks:
            logger.warning(
                "All documents loaded but produced zero chunks. "
                "Check that files are not empty or image-only PDFs."
            )
            if force_rebuild:
                self._clear_index()
            return 0

        logger.info("Total chunks to embed: %d", len(all_chunks))

        # ── Build FAISS index ─────────────────────────────────────────────
        from langchain_community.vectorstores import FAISS  # lazy
        from langchain_core.documents import Document  # lazy

        lc_docs = [
            Document(
                page_content=c["text"],
                metadata={
                    "source": c["source"],
                    "page": c["page"],
                    "chunk_index": c["chunk_index"],
                },
            )
            for c in all_chunks
        ]

        embeddings = self._get_embeddings()
        logger.info("Embedding %d chunks — this may take a moment …", len(lc_docs))

        # Embed in small batches rather than one FAISS.from_documents() call
        # over the entire document set. Embedding all chunks at once holds
        # every chunk's text AND its resulting vector in memory
        # simultaneously (on top of the already-loaded ONNX runtime,
        # Streamlit, and LangChain baseline) — on memory-constrained hosts
        # (e.g. Render's 512MB starter plan) this was observed to exceed
        # the limit and crash the process mid-ingest. Batching keeps peak
        # memory roughly constant regardless of knowledge-base size, at
        # the cost of a few extra FAISS merge calls (index build still
        # only touches disk once, at the end, via the atomic swap below).
        batch_size = EMBEDDING_BATCH_SIZE
        new_store = None
        for start in range(0, len(lc_docs), batch_size):
            batch = lc_docs[start: start + batch_size]
            if new_store is None:
                new_store = FAISS.from_documents(batch, embeddings)
            else:
                new_store.add_documents(batch)
            logger.debug(
                "  embedded %d/%d chunks",
                min(start + batch_size, len(lc_docs)),
                len(lc_docs),
            )
            gc.collect()

        # ── Atomic write (temp-dir then swap) ─────────────────────────────
        self._vectorstore_path.mkdir(parents=True, exist_ok=True)
        tmp_dir = Path(
            tempfile.mkdtemp(
                prefix="edupilot_vs_tmp_",
                dir=self._vectorstore_path.parent,
            )
        )
        try:
            new_store.save_local(str(tmp_dir))
            self._write_meta()  # write meta.json into final location

            # Move each index file from tmp into place
            for filename in (INDEX_FILENAME, PKL_FILENAME):
                src = tmp_dir / filename
                dst = self._vectorstore_path / filename
                if src.exists():
                    shutil.move(str(src), str(dst))
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

        self._store = new_store
        logger.info(
            "FAISS index built and saved to '%s' (%d chunks).",
            self._vectorstore_path,
            len(all_chunks),
        )
        return len(all_chunks)

    def retrieve(
        self, query: str, top_k: int = 5
    ) -> List[Tuple[str, SourceAttribution]]:
        """Retrieve the top-k most relevant chunks for *query*.

        Lazily loads or builds the FAISS index on the first call so callers
        never need to call :meth:`load_or_create` manually.

        Args:
            query: The search query string.
            top_k: Number of chunks to return (default 5).

        Returns:
            List of ``(chunk_text, SourceAttribution)`` tuples ordered by
            relevance score descending. Returns an empty list (never raises)
            when the knowledge base is empty or not yet indexed.
        """
        if not query or not query.strip():
            logger.warning("retrieve() called with empty query; returning [].")
            return []

        # Lazy load / build
        if self._store is None:
            if not self._index_exists() and not self._knowledge_dir_has_docs():
                logger.warning(
                    "Knowledge base is empty (no documents in '%s' and no "
                    "persisted index). Returning empty retrieval results.",
                    self._knowledge_dir,
                )
                return []
            self.load_or_create()

        if self._store is None:
            # load_or_create() found nothing to index
            logger.warning(
                "Index unavailable after load attempt; returning []."
            )
            return []

        try:
            docs_scores = self._store.similarity_search_with_score(
                query, k=top_k
            )
        except Exception as exc:
            logger.error("FAISS similarity search failed: %s", exc)
            return []

        results: List[Tuple[str, SourceAttribution]] = []
        for doc, raw_score in docs_scores:
            meta = doc.metadata
            # FAISS returns L2 distance; convert to a 0-1 similarity proxy
            similarity = float(1.0 / (1.0 + raw_score))
            attribution = SourceAttribution(
                document_name=meta.get("source", "unknown"),
                page_number=meta.get("page") or None,
                relevance_score=round(similarity, 4),
                excerpt=doc.page_content[:200],
            )
            results.append((doc.page_content, attribution))

        logger.info(
            "Retrieved %d chunks for query: '%s…'",
            len(results),
            query[:60],
        )
        return results

    # ------------------------------------------------------------------
    # Internal utilities
    # ------------------------------------------------------------------

    def _knowledge_dir_has_docs(self) -> bool:
        """Return True if the knowledge directory contains at least one supported file.

        Returns:
            bool: True when a supported document exists in ``knowledge/``.
        """
        if not self._knowledge_dir.exists():
            return False
        return any(
            p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
            for p in self._knowledge_dir.iterdir()
        )


# ---------------------------------------------------------------------------
# CLI entry point — for local verification only
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Syllabus CO table extraction
# ---------------------------------------------------------------------------
# Separate from the RAG text-chunking pipeline above: this reads a syllabus/
# scheme PDF's actual table structure (via pdfplumber, not plain-text
# chunking) to pull out the official, faculty-approved Course Outcomes for a
# given course code — e.g. "CO1: Understand the Business Intelligence,
# Analytics and Decision Support system" — so generated papers can print the
# institution's real CO wording instead of the LLM's own guess.

_SYLLABUS_CO_CACHE: dict = {}
"""Cache keyed by (pdf_path, mtime, course_code) -> extracted CO rows, so a
syllabus PDF already scanned once (potentially 50+ pages) isn't re-parsed on
every single assessment generation."""


def _find_co_table_in_pdf(
    pdf_path, course_code: str
) -> List[Tuple[str, str, str]]:
    """Extract the CO table for *course_code* from one syllabus PDF.

    Locates the page containing a "Course Code : <course_code>" line, then
    scans that page and the next few for a table whose header row contains
    both "CO" and "Course Outcomes" columns (the institutional syllabus
    format — see the "RBT Level" / "RBT Level Indicator" columns that
    typically follow). Tolerates two-letter and multi-letter RBT indicator
    formats (e.g. "L2" or "Ap") since both appear across different course
    pages in practice.

    Args:
        pdf_path: Path to the syllabus/scheme PDF.
        course_code: Course code to search for (e.g. "BCS703A").

    Returns:
        List of ``(co_code, description, rbt_level)`` tuples in document
        order. Empty list if the course code or its CO table isn't found.
    """
    import pdfplumber  # lazy import — only needed for this feature

    code = course_code.strip().upper()
    if not code:
        return []

    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            start_idx = None
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                if re.search(
                    rf"Course\s*Code\s*:?\s*{re.escape(code)}\b",
                    text, re.IGNORECASE,
                ):
                    start_idx = i
                    break

            if start_idx is None:
                return []

            for i in range(start_idx, min(start_idx + 6, len(pdf.pages))):
                page = pdf.pages[i]
                text = page.extract_text() or ""
                if "Course Outcomes" not in text:
                    continue

                for table in page.extract_tables():
                    if not table or not table[0]:
                        continue
                    header = [(c or "").strip().lower() for c in table[0]]
                    has_co_col = any(h == "co" or h.startswith("co ") for h in header)
                    has_desc_col = any("course outcomes" in h for h in header)
                    if not (has_co_col and has_desc_col):
                        continue

                    rows: List[Tuple[str, str, str]] = []
                    for row in table[1:]:
                        if not row or not row[0]:
                            continue
                        co_code = (row[0] or "").strip()
                        if not re.match(r"^CO\d+$", co_code, re.IGNORECASE):
                            continue
                        desc = (row[1] or "").strip().replace("\n", " ")
                        rbt = (row[2] or "").strip().replace("\n", " ") if len(row) > 2 else ""
                        if desc:
                            rows.append((co_code.upper(), desc, rbt))
                    if rows:
                        return rows
        return []
    except Exception as exc:  # noqa: BLE001 — never let a bad PDF break generation
        logger.warning(
            "_find_co_table_in_pdf(): failed to parse %s for %r: %s",
            pdf_path, course_code, exc,
        )
        return []


def extract_syllabus_co_table(course_code: str) -> List[Tuple[str, str, str]]:
    """Find the official CO table for *course_code* across the knowledge base.

    Scans every PDF in ``config.knowledge_dir`` (not a hardcoded filename —
    whichever scheme/syllabus document the faculty has uploaded, under
    whatever name) for a "Course Code : <course_code>" match, returning the
    first CO table found. Results are cached per (file, mtime, course code)
    so repeated generations for the same course don't re-scan a large
    syllabus PDF every time.

    Args:
        course_code: Course code to look up (e.g. "BCS703A"). Case-insensitive.

    Returns:
        List of ``(co_code, description, rbt_level)`` tuples, empty if no
        matching course code / CO table is found anywhere in the knowledge
        base (e.g. the syllabus hasn't been uploaded, or this course isn't
        in it) — callers should treat this as "no official CO table
        available" and fall back gracefully, never raise.
    """
    code = (course_code or "").strip().upper()
    if not code:
        return []

    knowledge_dir = config.knowledge_dir
    if not knowledge_dir.exists():
        return []

    for pdf_path in sorted(knowledge_dir.glob("*.pdf")):
        try:
            mtime = pdf_path.stat().st_mtime
        except OSError:
            continue
        cache_key = (str(pdf_path), mtime, code)
        if cache_key in _SYLLABUS_CO_CACHE:
            cached = _SYLLABUS_CO_CACHE[cache_key]
            if cached:
                return cached
            continue

        rows = _find_co_table_in_pdf(pdf_path, code)
        _SYLLABUS_CO_CACHE[cache_key] = rows
        if rows:
            logger.info(
                "extract_syllabus_co_table(): found %d COs for %r in %s",
                len(rows), code, pdf_path.name,
            )
            return rows

    return []


# ---------------------------------------------------------------------------
# Content diagram extraction (figures, circuits, network diagrams)
# ---------------------------------------------------------------------------
# Separate again from both the RAG chunking pipeline and the syllabus CO
# parser above: this pulls genuine content images (technical figures,
# circuit diagrams, charts) out of knowledge-base PDFs — filtering out
# decorative template graphics (repeated banners, logos) — and catalogs
# them with whatever real page text is available, so a generated question
# can reference a real diagram instead of describing one from scratch.
#
# Text-only matching (no OCR): a diagram is only findable if its slide/page
# has actual extractable text nearby (a title, a caption). Diagrams on
# fully-flattened image slides (the whole slide is one scanned picture,
# common for textbook-figure screenshots) have no text to match against
# and simply stay uncatalogued-by-topic — see the conversation with the
# person that scoped this decision. A worthwhile future upgrade (OCR
# fallback) is intentionally NOT included here.

_DIAGRAM_DECORATIVE_THRESHOLD_FRACTION = 0.15
"""An image dimension repeating on at least this fraction of a document's
pages is treated as decorative template art (banners, logos) and excluded."""

_DIAGRAM_MIN_AREA_PX = 80 * 80
"""Images smaller than this (in source pixels) are treated as bullet-point
icons / decorative glyphs, not content diagrams."""

_DIAGRAM_MAX_ASPECT_RATIO = 3.0
"""Images more elongated than this (width:height or height:width) are
treated as decorative ribbon/banner graphics, not content diagrams —
real technical figures observed in practice are much closer to square."""

_DIAGRAM_MANIFEST_FILENAME = "manifest.json"

_STOPWORDS = {
    "the", "a", "an", "of", "and", "or", "to", "in", "on", "for", "with",
    "is", "are", "was", "were", "be", "this", "that", "at", "by", "as",
    "it", "its", "from", "using", "use", "used", "how", "what", "which",
}


def _find_decorative_image_dims(pdf) -> set:
    """Return the set of (width, height) pairs that repeat across a
    document's pages often enough to be template/decorative art."""
    n_pages = len(pdf.pages)
    if n_pages == 0:
        return set()
    dim_counter: dict = {}
    for page in pdf.pages:
        for img in page.images:
            dim = (round(img["width"]), round(img["height"]))
            dim_counter[dim] = dim_counter.get(dim, 0) + 1
    return {
        dim for dim, count in dim_counter.items()
        if count / n_pages >= _DIAGRAM_DECORATIVE_THRESHOLD_FRACTION
    }


def extract_content_diagrams(pdf_path: Path, source_name: str) -> List[dict]:
    """Extract genuine content diagrams from one PDF, filtering out noise.

    Two filters, applied together:
      1. Recurring-dimension filter — an image size appearing on >=15% of
         the document's pages is a repeated template element (banner,
         logo), not content.
      2. Aspect-ratio filter — images more elongated than 3:1 are ribbon/
         divider graphics, not technical figures (verified against real
         decorative banners: 8.46:1 and 4.91:1, vs real circuit/algorithm
         figures at ~1:1).

    Saves each qualifying image as a cropped PNG under
    ``config.diagrams_dir / <source_name>/``, and returns metadata for
    each (no captions are OCR'd — see module docstring above).

    Args:
        pdf_path: Path to the source PDF.
        source_name: Stem used for the output subfolder and filenames
            (typically the PDF's filename without extension).

    Returns:
        List of dicts, one per extracted diagram:
        ``{"id", "source_file", "page", "image_path", "context_text"}``.
        Empty list if the PDF can't be opened or has no qualifying images.
    """
    import pdfplumber  # lazy import — only needed for this feature

    out_dir = config.diagrams_dir / source_name
    out_dir.mkdir(parents=True, exist_ok=True)

    manifest: List[dict] = []
    try:
        with pdfplumber.open(str(pdf_path)) as pdf:
            decorative_dims = _find_decorative_image_dims(pdf)
            idx = 0
            for page_num, page in enumerate(pdf.pages, start=1):
                if page_num == 1:
                    # Cover/title pages almost universally carry the
                    # institution logo/crest, not a content diagram — and
                    # a one-off cover logo can slip past both filters
                    # above (moderate aspect ratio, and it may not repeat
                    # often enough within THIS document to trip the
                    # recurring-dimension check, even though it's
                    # unmistakably decorative). Observed in practice: a
                    # college crest on a syllabus PDF's cover page got
                    # catalogued and attached to a generated question.
                    continue
                page_text = (page.extract_text() or "").strip()
                for img in page.images:
                    dim = (round(img["width"]), round(img["height"]))
                    if dim in decorative_dims:
                        continue
                    w, h = img["width"], img["height"]
                    if w <= 0 or h <= 0 or w * h < _DIAGRAM_MIN_AREA_PX:
                        continue
                    if max(w / h, h / w) > _DIAGRAM_MAX_ASPECT_RATIO:
                        continue
                    try:
                        bbox = (img["x0"], img["top"], img["x1"], img["bottom"])
                        cropped_img = page.crop(bbox).to_image(resolution=150)
                    except Exception as exc:  # noqa: BLE001
                        logger.debug(
                            "extract_content_diagrams(): skipped an image "
                            "on %s page %d: %s", source_name, page_num, exc,
                        )
                        continue
                    idx += 1
                    fname = f"{source_name}_p{page_num}_{idx}.png"
                    fpath = out_dir / fname
                    cropped_img.save(str(fpath))
                    manifest.append({
                        "id": f"{source_name}_p{page_num}_{idx}",
                        "source_file": source_name,
                        "page": page_num,
                        "image_path": str(fpath),
                        "context_text": page_text[:400],
                    })
    except Exception as exc:  # noqa: BLE001 — never let a bad PDF break ingest
        logger.warning(
            "extract_content_diagrams(): failed to process %s: %s",
            pdf_path, exc,
        )
        return []

    logger.info(
        "extract_content_diagrams(): extracted %d content diagram(s) from %s",
        len(manifest), source_name,
    )
    return manifest


def catalog_diagrams_for_knowledge_base() -> int:
    """Extract and catalog content diagrams from every PDF in the knowledge base.

    Scans ``config.knowledge_dir`` for PDFs, runs
    :func:`extract_content_diagrams` on each, and writes a combined
    manifest to ``config.diagrams_dir / manifest.json``. Safe to call
    repeatedly — re-running replaces the manifest and re-extracts (no
    incremental diffing, since this is expected to run alongside the
    existing full-rebuild knowledge base ingestion, not on every request).

    Returns:
        int: Total number of content diagrams catalogued across all PDFs.
    """
    knowledge_dir = config.knowledge_dir
    if not knowledge_dir.exists():
        return 0

    all_entries: List[dict] = []
    for pdf_path in sorted(knowledge_dir.glob("*.pdf")):
        source_name = re.sub(r"[^\w\-]", "_", pdf_path.stem)
        entries = extract_content_diagrams(pdf_path, source_name)
        all_entries.extend(entries)

    config.diagrams_dir.mkdir(parents=True, exist_ok=True)
    manifest_path = config.diagrams_dir / _DIAGRAM_MANIFEST_FILENAME
    try:
        manifest_path.write_text(json.dumps(all_entries, indent=2))
    except OSError as exc:
        logger.warning(
            "catalog_diagrams_for_knowledge_base(): failed to write "
            "manifest: %s", exc,
        )

    logger.info(
        "catalog_diagrams_for_knowledge_base(): catalogued %d diagram(s) "
        "across %d PDF(s)",
        len(all_entries),
        len(list(knowledge_dir.glob("*.pdf"))),
    )
    return len(all_entries)


def _load_diagram_manifest() -> List[dict]:
    """Load the cataloged-diagram manifest from disk, if it exists."""
    manifest_path = config.diagrams_dir / _DIAGRAM_MANIFEST_FILENAME
    if not manifest_path.exists():
        return []
    try:
        return json.loads(manifest_path.read_text())
    except (OSError, json.JSONDecodeError) as exc:
        logger.warning("_load_diagram_manifest(): failed to read: %s", exc)
        return []


def _keyword_overlap_score(query: str, candidate: str) -> int:
    """Count overlapping non-stopword keywords between two texts."""
    def _words(s: str) -> set:
        return {
            w for w in re.findall(r"[a-z0-9]+", s.lower())
            if len(w) > 2 and w not in _STOPWORDS
        }
    return len(_words(query) & _words(candidate))


def find_relevant_diagrams(query_text: str, top_k: int = 1) -> List[dict]:
    """Find catalogued diagrams whose page text best matches *query_text*.

    Simple keyword-overlap scoring (no embeddings/LLM call — fast,
    deterministic, and easy to reason about for a first version). Only
    diagrams with at least one overlapping keyword are returned, so a
    topic with no genuinely related diagram gets nothing back rather
    than a weakly-related best-effort guess.

    Args:
        query_text: The question text or topic to match against (e.g.
            a generated question's text, or the faculty's topic string).
        top_k: Maximum number of diagrams to return.

    Returns:
        List of manifest entries (see :func:`extract_content_diagrams`),
        best match first. Empty list if no catalog exists yet or nothing
        scores above zero overlap.
    """
    manifest = _load_diagram_manifest()
    if not manifest or not query_text:
        return []

    scored = [
        (entry, _keyword_overlap_score(query_text, entry.get("context_text", "")))
        for entry in manifest
    ]
    scored = [(e, s) for e, s in scored if s > 0]
    scored.sort(key=lambda pair: pair[1], reverse=True)
    return [e for e, _ in scored[:top_k]]


def attach_diagrams_to_questions(questions: List[Question]) -> int:
    """Attach a matching catalogued diagram to each question, where one exists.

    Loads the diagram manifest once (not per-question) and matches each
    question's own text against it via :func:`find_relevant_diagrams`'s
    scoring logic. Mutates ``question.diagram_path`` in place for any
    question with a positive-scoring match; leaves it empty otherwise —
    most questions won't have a relevant diagram, and that's expected,
    not an error.

    Args:
        questions: Questions to check and mutate in place.

    Returns:
        int: Number of questions that got a diagram attached.
    """
    manifest = _load_diagram_manifest()
    if not manifest:
        return 0

    attached = 0
    for q in questions:
        if getattr(q, "diagram_path", ""):
            continue  # already set (e.g. by a previous pass)
        scored = [
            (entry, _keyword_overlap_score(q.question_text or "", entry.get("context_text", "")))
            for entry in manifest
        ]
        scored = [(e, s) for e, s in scored if s > 0]
        if not scored:
            continue
        scored.sort(key=lambda pair: pair[1], reverse=True)
        best_entry, _ = scored[0]
        q.diagram_path = best_entry["image_path"]
        attached += 1

    if attached:
        logger.info(
            "attach_diagrams_to_questions(): attached diagrams to %d/%d questions",
            attached, len(questions),
        )
    return attached


def _cli_main() -> None:
    """Command-line verification helper.

    Usage::

        cd artifacts/edupilot
        python rag.py               # normal ingest + retrieve
        python rag.py --rebuild     # force full rebuild
        python rag.py --empty       # test empty-KB path
    """
    import argparse

    parser = argparse.ArgumentParser(description="EduPilot RAG CLI verifier")
    parser.add_argument(
        "--rebuild", action="store_true", help="Force index rebuild"
    )
    parser.add_argument(
        "--empty", action="store_true",
        help="Test empty knowledge base path (temporarily renames knowledge/)"
    )
    args = parser.parse_args()

    rag = RAGModule()

    if args.empty:
        print("\n=== Empty KB path test ===")
        # Temporarily point to a non-existent dir
        rag._knowledge_dir = Path("/tmp/edupilot_empty_test_kb")
        rag._vectorstore_path = Path("/tmp/edupilot_empty_test_vs")
        count = rag.ingest()
        print(f"ingest() returned: {count}  (expected 0)")
        results = rag.retrieve("any query")
        print(f"retrieve() returned: {results}  (expected [])")
        return

    print(f"\n=== EduPilot RAG CLI Verifier ===")
    print(f"Knowledge dir : {rag._knowledge_dir}")
    print(f"Vectorstore   : {rag._vectorstore_path}")
    print(f"Embedding     : {rag._embedding_model_name}")
    print()

    # Ingest
    count = rag.ingest(force_rebuild=args.rebuild)
    print(f"\n✓ ingest() → {count} chunks indexed")

    if count == 0:
        print("  (No documents found — add files to knowledge/ and retry)")
        return

    # Retrieve
    test_query = "data structures and algorithms"
    print(f"\n=== Retrieval test: '{test_query}' (top_k=3) ===")
    results = rag.retrieve(test_query, top_k=3)
    print(f"✓ retrieve() → {len(results)} results\n")
    for i, (text, attr) in enumerate(results, 1):
        print(f"  [{i}] source={attr.document_name!r}  page={attr.page_number}"
              f"  score={attr.relevance_score:.4f}")
        print(f"      excerpt: {text[:120].replace(chr(10), ' ')!r}")
        print()

    # format_rag_context
    ctx = format_rag_context(results)
    print("=== format_rag_context() output ===")
    print(ctx)

    # Reload test — simulates a fresh process
    print("\n=== Reload from disk test ===")
    rag2 = RAGModule()
    results2 = rag2.retrieve(test_query, top_k=2)
    print(f"✓ Fresh RAGModule.retrieve() → {len(results2)} results (from disk)")


if __name__ == "__main__":
    _cli_main()