"""
EduPilot AI Faculty Assistant
Module: downloads.py
Version: 4.1.0
Author: EduPilot Team
Purpose: Professional document generation engine. Exports an Assessment to
         Markdown, Microsoft Word (.docx via python-docx), and PDF (via
         ReportLab). All three formats share a common parsed-structure
         intermediate representation so that documents are built from typed
         data rather than raw LLM text strings. Includes university-grade
         header/footer, page numbers, Bloom / CO / difficulty metadata,
         question numbering, and an answer-key section. Gracefully handles
         optional/missing fields so a partially-generated assessment never
         produces a broken document.
"""

from __future__ import annotations

import io
import re
import textwrap
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    HRFlowable,
    Image,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.flowables import KeepTogether

from docx import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from logging_utils import get_logger
from models import Assessment, AssessmentType, Question, split_sizes_for_pairing

logger = get_logger(__name__)

# ---------------------------------------------------------------------------
# Constants — University branding defaults
# ---------------------------------------------------------------------------

_UNIVERSITY_NAME = "Dayananda Sagar Academy of Technology and Management"
_UNIVERSITY_AFFILIATION = "(Autonomous Institute under VTU)"
_UNIVERSITY_TAGLINE = "Department of Academic Excellence"
_UNIVERSITY_ACCREDITATION_LINES = [
    "Affiliated to VTU",
    "Approved by AICTE",
    "Accredited by NAAC with A+ Grade",
    "4 Programs Accredited by NBA (CSE, ISE, ECE, ME)",
]
# Words highlighted in red within the accreditation lines above, matching
# the official paper's styling exactly.
_ACCREDITATION_HIGHLIGHT_WORDS = (
    "VTU", "AICTE", "NAAC", "A+", "NBA",
)
_ACCREDITATION_HIGHLIGHT_RE = re.compile(
    "(" + "|".join(
        re.escape(w) for w in
        sorted(_ACCREDITATION_HIGHLIGHT_WORDS, key=len, reverse=True)
    ) + ")"
)
_ACCREDITATION_HIGHLIGHT_COLOR = "CC0000"  # red, matches the official paper
# Institution logo (left) and IQAC accreditation badge (right) for the exam
# paper header. Drop the actual image files at these paths — both are
# optional: if a file is missing, the header simply renders without that
# image (text-only fallback) rather than failing, so this works today and
# picks up real artwork the moment it's added.
_ASSETS_DIR = Path(__file__).resolve().parent / "assets"
_INSTITUTION_LOGO_PATH = _ASSETS_DIR / "institution_logo.png"
_IQAC_BADGE_PATH = _ASSETS_DIR / "iqac_badge.png"


def _department_heading(department: str) -> str:
    """Return a "Department of X" heading without doubling the prefix.

    ``ParsedAssessment.department`` falls back to a generic tagline
    ("Department of Academic Excellence") when no real department was
    set on the assessment — that fallback already reads naturally on
    its own, so prefixing it again would print "Department of
    Department of Academic Excellence". Real department names (e.g.
    "Computer Science and Engineering") don't start with "Department",
    so they still get the prefix as expected.
    """
    dept = (department or "").strip()
    if dept.lower().startswith("department"):
        return dept
    return f"Department of {dept}" if dept else "Department"
_DATE_FMT = "%d %B %Y"
_DATE_TIME_FMT = "%d %B %Y, %I:%M %p"


# ===========================================================================
# Shared intermediate representation
# ===========================================================================


@dataclass
class ParsedQuestion:
    """Flat, rendering-friendly view of a single Question.

    All fields are strings so formatters need no further type handling.
    """

    number: int
    question_id: str
    text: str
    question_type: str
    bloom_level: str
    difficulty: str
    co_mapping: str          # comma-joined list, e.g. "CO1, CO3"
    marks: str               # stringified integer
    answer_key: str
    notes: str
    options: List[str] = field(default_factory=list)  # MCQ answer options
    topic: str = ""  # syllabus topic/module this question targets
    blueprint_group: str = ""  # "A"/"B" when a custom marks blueprint was used
    case_background: str = ""  # Consultancy Case Study: scenario/context text


@dataclass
class ParsedAssessment:
    """Structured, renderer-neutral view of an Assessment.

    Built once by :class:`AssessmentParser` and consumed by all three
    exporters so that parsing logic lives in exactly one place.
    """

    # Header fields
    university: str
    department: str
    title: str
    course_code: str
    course_name: str
    assessment_type: str
    semester: str
    duration: str           # "90 minutes"
    total_marks: str        # "100 marks"
    faculty_name: str
    date_generated: str
    datetime_generated: str
    test_date: str
    instructions: str
    batch: str
    teaching_department: str
    academic_year: str

    # Body
    questions: List[ParsedQuestion] = field(default_factory=list)
    generation_notes: str = ""

    # Derived helpers
    has_answer_key: bool = False
    has_instructions: bool = False
    has_generation_notes: bool = False


# ---------------------------------------------------------------------------
# Parser
# ---------------------------------------------------------------------------


class AssessmentParser:
    """Convert an :class:`~models.Assessment` into a :class:`ParsedAssessment`.

    The parser is the single source of truth for flattening the typed
    Assessment model into plain strings that all exporters can consume
    without re-implementing the same null-checks.
    """

    # Strips leading/trailing whitespace and normalises internal whitespace runs
    _WS_RE = re.compile(r"\s+")

    @classmethod
    def parse(cls, assessment: Assessment) -> ParsedAssessment:
        """Parse an Assessment into a rendering-ready ParsedAssessment.

        Args:
            assessment: The typed Assessment object to render.

        Returns:
            A fully populated :class:`ParsedAssessment` ready for export.
        """
        meta = assessment.metadata

        # ── Header ──────────────────────────────────────────────────────────
        department = cls._safe(meta.department) or _UNIVERSITY_TAGLINE
        faculty_name = cls._safe(meta.faculty_name) or "Faculty"
        semester = cls._safe(meta.semester) or "Current Semester"
        instructions = cls._safe(meta.instructions)

        duration = (
            f"{meta.duration_minutes} minutes"
            if meta.duration_minutes > 0
            else "As specified"
        )
        total_marks_computed = assessment.total_marks or meta.total_marks
        total_marks = (
            f"{total_marks_computed} marks"
            if total_marks_computed > 0
            else "As specified"
        )

        # ── Questions ───────────────────────────────────────────────────────
        parsed_questions: List[ParsedQuestion] = []
        has_any_answer_key = False

        for idx, q in enumerate(assessment.questions, start=1):
            co_str = ", ".join(q.co_mapping) if q.co_mapping else "—"
            answer_key_text = cls._safe(q.answer_key)
            if answer_key_text:
                has_any_answer_key = True

            parsed_questions.append(
                ParsedQuestion(
                    number=idx,
                    question_id=cls._safe(q.question_id) or f"Q{idx}",
                    text=cls._clean(q.question_text),
                    question_type=cls._safe(q.question_type) or "Short Answer",
                    bloom_level=q.bloom_level.value
                    if hasattr(q.bloom_level, "value")
                    else str(q.bloom_level),
                    difficulty=q.difficulty.value
                    if hasattr(q.difficulty, "value")
                    else str(q.difficulty),
                    co_mapping=co_str,
                    marks=str(q.marks),
                    answer_key=answer_key_text,
                    notes=cls._safe(q.notes),
                    options=[
                        cls._safe(o) for o in getattr(q, "options", []) or []
                        if cls._safe(o)
                    ],
                    topic=cls._safe(getattr(q, "topic", "")),
                    blueprint_group=cls._safe(
                        getattr(q, "blueprint_group", "")
                    ),
                    case_background=cls._safe(
                        getattr(q, "case_background", "")
                    ),
                )
            )

        generation_notes = cls._safe(assessment.generation_notes)

        return ParsedAssessment(
            university=_UNIVERSITY_NAME,
            department=department,
            title=cls._safe(meta.title) or "Assessment",
            course_code=cls._safe(meta.course_code) or "",
            course_name=cls._safe(meta.course_name) or "",
            assessment_type=meta.assessment_type.value
            if hasattr(meta.assessment_type, "value")
            else str(meta.assessment_type),
            semester=semester,
            duration=duration,
            total_marks=total_marks,
            faculty_name=faculty_name,
            date_generated=datetime.now().strftime(_DATE_FMT),
            datetime_generated=datetime.now().strftime(_DATE_TIME_FMT),
            test_date=cls._safe(getattr(meta, "test_date", "")),
            instructions=instructions,
            batch=cls._safe(getattr(meta, "batch", "")),
            teaching_department=(
                cls._safe(getattr(meta, "teaching_department", "")) or department
            ),
            academic_year=cls._safe(getattr(meta, "academic_year", "")),
            questions=parsed_questions,
            generation_notes=generation_notes,
            has_answer_key=has_any_answer_key,
            has_instructions=bool(instructions),
            has_generation_notes=bool(generation_notes),
        )

    # ── Helpers ──────────────────────────────────────────────────────────────

    @staticmethod
    def _safe(value: object) -> str:
        """Coerce *value* to a stripped string; return empty string for None."""
        if value is None:
            return ""
        return str(value).strip()

    @classmethod
    def _clean(cls, text: str) -> str:
        """Strip, collapse internal whitespace runs, and return the result."""
        return cls._WS_RE.sub(" ", text.strip()) if text else ""


# ===========================================================================
# VTU-style Semester Examination paper layout
# ===========================================================================
#
# The Semester Examination export follows the standard VTU (Visvesvaraya
# Technological University) semester-end exam paper convention: questions
# are grouped into numbered Modules (one per syllabus topic), each Module
# offers exactly one pair of alternative full questions ("Q1 OR Q2"), and
# each full question is broken into lettered sub-parts (a, b, c...) — each
# sub-part being one AI-generated question. This section builds that
# grouped structure from the flat ParsedQuestion list; WordExporter and
# PDFExporter both consume it identically.


@dataclass
class VTUSubQuestion:
    """One lettered sub-part (a, b, c…) of a VTU-style full question."""

    letter: str
    question: ParsedQuestion


@dataclass
class VTUQuestionGroup:
    """One full numbered question (e.g. "Q1"), made up of 1+ sub-parts."""

    q_number: int
    subparts: List[VTUSubQuestion] = field(default_factory=list)

    @property
    def total_marks(self) -> int:
        """Sum of marks across this question's sub-parts."""
        total = 0
        for sp in self.subparts:
            try:
                total += int(sp.question.marks)
            except (TypeError, ValueError):
                continue
        return total


@dataclass
class VTUModule:
    """One syllabus Module, containing one or more OR-alternative pairs."""

    module_number: int
    label: str
    pairs: List[List[VTUQuestionGroup]] = field(default_factory=list)
    """Each entry is a list of 1 or 2 VTUQuestionGroup — 2 means a genuine
    "Q(n) OR Q(n+1)" alternative pair; 1 means a lone question with no
    alternative (only occurs when a Module has just a single question)."""


def build_vtu_paper_layout(
    questions: List[ParsedQuestion], max_subparts: int = 3
) -> List[VTUModule]:
    """Group a flat question list into VTU-style Modules with OR-pairs.

    Questions are grouped by their ``topic`` field, in first-seen order
    (which — because batched generation processes topics in the plan's
    original order — reliably matches the syllabus's module order even
    though no explicit module number is tracked elsewhere in the app).
    Questions with no topic (e.g. very old saved runs predating the
    ``topic`` field) are grouped under a single "General" module.

    Args:
        questions: Flat, already-parsed question list in original order.
        max_subparts: Maximum lettered sub-parts per full question.

    Returns:
        List of VTUModule, one per distinct topic, in first-seen order.
    """
    groups_by_topic: dict = {}
    topic_order: List[str] = []
    for q in questions:
        key = q.topic.strip() if q.topic and q.topic.strip() else "General"
        if key not in groups_by_topic:
            groups_by_topic[key] = []
            topic_order.append(key)
        groups_by_topic[key].append(q)

    letters = "abcdefgh"
    modules: List[VTUModule] = []
    q_counter = 1
    for m_idx, topic in enumerate(topic_order, start=1):
        topic_questions = groups_by_topic[topic]

        # When every question in this topic carries a blueprint_group
        # ("A"/"B" — set when a faculty custom marks blueprint was used
        # for this topic), split precisely along that boundary instead of
        # guessing an even split. This preserves the faculty's exact
        # intended sub-question structure for each OR-alternative.
        has_full_blueprint_tagging = bool(topic_questions) and all(
            q.blueprint_group in ("A", "B") for q in topic_questions
        )
        if has_full_blueprint_tagging:
            chunks = [
                [q for q in topic_questions if q.blueprint_group == "A"],
                [q for q in topic_questions if q.blueprint_group == "B"],
            ]
            chunks = [c for c in chunks if c]  # drop an empty side, if any
        else:
            sizes = split_sizes_for_pairing(len(topic_questions), max_subparts)
            chunks = []
            idx = 0
            for size in sizes:
                chunks.append(topic_questions[idx: idx + size])
                idx += size

        groups: List[VTUQuestionGroup] = []
        for chunk in chunks:
            subparts = [
                VTUSubQuestion(letter=letters[i], question=cq)
                for i, cq in enumerate(chunk)
            ]
            groups.append(VTUQuestionGroup(q_number=q_counter, subparts=subparts))
            q_counter += 1
        pairs = [groups[i: i + 2] for i in range(0, len(groups), 2)]
        modules.append(VTUModule(module_number=m_idx, label=topic, pairs=pairs))
    return modules


def _mirror_pair_co_bloom(pair: List["VTUQuestionGroup"]) -> None:
    """Mirror the first side's CO/Bloom pattern onto the second side.

    Q(n) and Q(n+1) are two alternative versions of the SAME question —
    a student answers only one — so if Q(n)'s sub-parts are CO1/CO2/CO1
    at Bloom levels L1/L3/L1, its OR-counterpart Q(n+1) should carry
    that identical pattern across its own a)/b)/c), even though the
    question text differs. This is a display-level safety net: agent.py
    already applies this at generation time for blueprinted assessments
    (see AssessmentAgent._apply_blueprint_marks), but export can be
    called on older saved runs generated before that fix existed, so
    this re-applies the same rule at render time regardless of when or
    how the assessment was generated. A no-op when the pair has only
    one side (no OR-alternative to mirror against).

    Args:
        pair: A ``module.pairs[i]`` entry — a list of 1 or 2
            VTUQuestionGroup. Mutates the second group's ParsedQuestion
            objects in place (safe: this is the export-time parsed copy,
            not the original stored Question objects).
    """
    if len(pair) != 2:
        return
    side_a, side_b = pair[0].subparts, pair[1].subparts
    for a_sp, b_sp in zip(side_a, side_b):
        b_sp.question.co_mapping = a_sp.question.co_mapping
        b_sp.question.bloom_level = a_sp.question.bloom_level


def _compute_co_coverage(modules: List["VTUModule"]) -> List[tuple]:
    """Compute marks-weighted CO coverage percentages for a VTU-style paper.

    Only the achievable side of each OR-pair (index 0 — the side actually
    counted toward the paper's max marks, matching the ``true_max_marks``
    calculation used elsewhere) is counted, since a student only ever
    answers one side per pair and _mirror_pair_co_bloom already keeps
    both sides' CO/Bloom pattern identical anyway. When a sub-part maps
    to multiple COs, its marks are attributed to the first (primary) CO
    only, so percentages sum to 100% rather than double-counting.

    Args:
        modules: Pre-built VTU module/OR-pair layout.

    Returns:
        List of ``(co_code, percentage)`` tuples for every CO that
        appears at least once, sorted by CO code (e.g. "CO1" before
        "CO2"). Empty list when there are no marks to attribute.
    """
    co_marks: dict = {}
    total_marks = 0
    for m in modules:
        for pair in m.pairs[:1]:
            for group in pair[:1]:
                for sp in group.subparts:
                    try:
                        marks = int(sp.question.marks)
                    except (TypeError, ValueError):
                        continue
                    total_marks += marks
                    co_str = (sp.question.co_mapping or "").strip()
                    primary_co = co_str.split(",")[0].strip() if co_str else ""
                    if not primary_co:
                        continue
                    co_marks[primary_co] = co_marks.get(primary_co, 0) + marks

    if total_marks == 0:
        return []

    def _co_sort_key(co: str):
        digits = "".join(ch for ch in co if ch.isdigit())
        return (0, int(digits)) if digits else (1, co)

    return [
        (co, round(marks * 100 / total_marks))
        for co, marks in sorted(co_marks.items(), key=lambda kv: _co_sort_key(kv[0]))
    ]


def build_fixed_pair_layout(
    questions: List[ParsedQuestion], pair_count: int = 5
) -> List[VTUModule]:
    """Group a flat question list into exactly ``pair_count`` OR-pairs.

    Two grouping strategies, chosen automatically:

    1. **Blueprint-aware** (preferred, used whenever every question in a
       topic carries a faculty-authored ``blueprint_group`` of "A" or
       "B" — see the Custom Sub-Question Marks field on the Generate
       form). Each topic becomes exactly one pair, split precisely along
       the A/B boundary the faculty specified, so multi-part questions
       (e.g. a), b), c) sub-parts with faculty-chosen marks) render
       exactly as authored. This is the same per-topic blueprint logic
       :func:`build_vtu_paper_layout` already uses for Semester
       Examination, reused here for Internal Assessment.
    2. **Flat fallback** (used when blueprint tagging is absent or only
       partial): splits the flat, ordered question list evenly across
       ``pair_count * 2`` groups regardless of topic. This exists
       because topic-based grouping alone is unreliable — if every
       question ends up tagged with the same (or an empty) topic, topic
       grouping collapses everything into a single giant pair instead of
       ``pair_count`` — so this guarantees the target pair count no
       matter how (or whether) topics were tagged upstream.

    Args:
        questions: Flat, already-parsed question list in original order.
        pair_count: Number of OR-pairs to produce (5 for a standard IAT).

    Returns:
        List of VTUModule, one per pair, each with an empty label (IAT
        papers render without "Module – N" divider rows — see
        ``show_module_labels=False`` in the exporters).
    """
    if not questions:
        return []

    # ── Strategy 1: blueprint-aware, grouped by topic ──────────────────
    groups_by_topic: dict = {}
    topic_order: List[str] = []
    for q in questions:
        key = q.topic.strip() if q.topic and q.topic.strip() else "General"
        if key not in groups_by_topic:
            groups_by_topic[key] = []
            topic_order.append(key)
        groups_by_topic[key].append(q)

    all_topics_blueprinted = all(
        bool(qs) and all(q.blueprint_group in ("A", "B") for q in qs)
        for qs in groups_by_topic.values()
    )

    letters = "abcdefgh"
    if all_topics_blueprinted and len(topic_order) >= 1:
        modules: List[VTUModule] = []
        q_counter = 1
        for m_idx, topic in enumerate(topic_order, start=1):
            topic_questions = groups_by_topic[topic]
            chunks = [
                [q for q in topic_questions if q.blueprint_group == "A"],
                [q for q in topic_questions if q.blueprint_group == "B"],
            ]
            chunks = [c for c in chunks if c]
            groups: List[VTUQuestionGroup] = []
            for chunk in chunks:
                subparts = [
                    VTUSubQuestion(letter=letters[i], question=cq)
                    for i, cq in enumerate(chunk)
                ]
                groups.append(
                    VTUQuestionGroup(q_number=q_counter, subparts=subparts)
                )
                q_counter += 1
            modules.append(
                VTUModule(module_number=m_idx, label="", pairs=[groups])
            )
        return modules

    # ── Strategy 2: flat fallback, ignores topic entirely ──────────────
    total = len(questions)
    group_count = min(pair_count * 2, total)
    if group_count == 0:
        return []

    base, remainder = divmod(total, group_count)
    chunks: List[List[ParsedQuestion]] = []
    idx = 0
    for i in range(group_count):
        size = base + (1 if i < remainder else 0)
        chunks.append(questions[idx: idx + size])
        idx += size

    groups: List[VTUQuestionGroup] = []
    for i, chunk in enumerate(chunks, start=1):
        subparts = [
            VTUSubQuestion(letter=letters[j], question=cq)
            for j, cq in enumerate(chunk)
        ]
        groups.append(VTUQuestionGroup(q_number=i, subparts=subparts))

    modules: List[VTUModule] = []
    for p in range(0, len(groups), 2):
        pair = groups[p: p + 2]
        modules.append(
            VTUModule(module_number=(p // 2) + 1, label="", pairs=[pair])
        )
    return modules


# ===========================================================================
# Markdown exporter
# ===========================================================================


class MarkdownExporter:
    """Export an :class:`~models.Assessment` to a clean Markdown string.

    The output is spec-compliant CommonMark with proper heading hierarchy,
    a metadata table, numbered questions, per-question metadata, and an
    answer-key section.
    """

    def export(self, assessment: Assessment) -> str:
        """Render the assessment as a Markdown document.

        Args:
            assessment: The typed Assessment to export.

        Returns:
            str: UTF-8 Markdown document string.

        Raises:
            ValueError: If *assessment* is None.
        """
        if assessment is None:
            raise ValueError("assessment must not be None")

        logger.info(
            "MarkdownExporter: exporting '%s'",
            assessment.metadata.title,
        )

        pa = AssessmentParser.parse(assessment)
        lines: List[str] = []

        # ── Document header ──────────────────────────────────────────────────
        lines += [
            f"# {pa.university}",
            f"### {pa.department}",
            "",
            f"## {pa.title}",
            "",
        ]

        # Metadata table
        meta_rows = [
            ("Course Code", pa.course_code or "—"),
            ("Course Name", pa.course_name or "—"),
            ("Assessment Type", pa.assessment_type),
            ("Semester", pa.semester),
            ("Duration", pa.duration),
            ("Total Marks", pa.total_marks),
            ("Faculty", pa.faculty_name),
            ("Test Date", pa.test_date or "—"),
            ("Generated On", pa.date_generated),
        ]
        lines += [
            "| Field | Details |",
            "|-------|---------|",
        ]
        for label, value in meta_rows:
            lines.append(f"| **{label}** | {value} |")
        lines.append("")

        # Instructions
        if pa.has_instructions:
            lines += [
                "---",
                "",
                "### Instructions",
                "",
                pa.instructions,
                "",
            ]

        # ── Questions ────────────────────────────────────────────────────────
        lines += [
            "---",
            "",
            "## Questions",
            "",
        ]

        for q in pa.questions:
            lines += [
                f"### {q.number}. {q.text}",
                "",
            ]
            if q.options:
                lines += [
                    f"- ( ) **{chr(65 + i)}.** {opt}"
                    for i, opt in enumerate(q.options)
                ] + [""]
            lines += [
                "| Attribute | Value |",
                "|-----------|-------|",
                f"| **Type** | {q.question_type} |",
                f"| **Bloom Level** | {q.bloom_level} |",
                f"| **Difficulty** | {q.difficulty} |",
                f"| **CO Mapping** | {q.co_mapping} |",
                f"| **Marks** | {q.marks} |",
                "",
            ]
            if q.notes:
                lines += [f"> **Note:** {q.notes}", ""]

        # ── Answer key ───────────────────────────────────────────────────────
        if pa.has_answer_key:
            lines += [
                "---",
                "",
                "## Answer Key",
                "",
                "> *For faculty use only — not to be distributed to students.*",
                "",
            ]
            for q in pa.questions:
                answer_text = q.answer_key if q.answer_key else "*Not provided*"
                lines += [
                    f"### {q.number}. {q.question_id}",
                    "",
                    answer_text,
                    "",
                ]

        # ── Generation notes ─────────────────────────────────────────────────
        if pa.has_generation_notes:
            lines += [
                "---",
                "",
                "## Agent Notes",
                "",
                pa.generation_notes,
                "",
            ]

        # ── Footer ───────────────────────────────────────────────────────────
        lines += [
            "---",
            "",
            f"*Generated by EduPilot AI Faculty Assistant · {pa.date_generated}*",
            "",
        ]

        md = "\n".join(lines)
        logger.info(
            "MarkdownExporter: produced %d characters", len(md)
        )
        return md


# ===========================================================================
# Word exporter
# ===========================================================================


class WordExporter:
    """Export an :class:`~models.Assessment` to a professional DOCX byte stream.

    Uses ``python-docx`` to construct a formatted Word document with:
    - University header with branding colours
    - Metadata table
    - Numbered questions with Bloom/CO/difficulty metadata
    - Answer-key section (page-break separated, faculty-only notice)
    - Automatic page numbers in the footer
    """

    # Colour palette (RGB tuples)
    _COLOUR_PRIMARY = RGBColor(0x1A, 0x37, 0x6C)   # deep navy
    _COLOUR_ACCENT = RGBColor(0xC8, 0x9B, 0x20)    # gold
    _COLOUR_LIGHT = RGBColor(0xF0, 0xF4, 0xF8)     # light blue-grey

    def export(self, assessment: Assessment) -> bytes:
        """Render the assessment as a DOCX byte stream.

        Args:
            assessment: The typed Assessment to export.

        Returns:
            bytes: Raw .docx bytes suitable for ``st.download_button``.

        Raises:
            ValueError: If *assessment* is None.
        """
        if assessment is None:
            raise ValueError("assessment must not be None")

        logger.info(
            "WordExporter: exporting '%s'", assessment.metadata.title
        )

        pa = AssessmentParser.parse(assessment)
        doc = DocxDocument()
        is_semester_exam = (
            assessment.metadata.assessment_type == AssessmentType.SEMESTER_EXAM
        )
        is_internal_assessment = (
            assessment.metadata.assessment_type == AssessmentType.INTERNAL
        )
        is_case_study = (
            assessment.metadata.assessment_type
            == AssessmentType.CONSULTANCY_CASE
        )

        self._configure_page(doc)
        self._add_footer(doc, pa)

        if is_semester_exam or is_internal_assessment:
            # Both Semester Examination and Internal Assessment use the
            # same OR-alternative paper structure. Semester Examination
            # groups by syllabus topic into Modules (build_vtu_paper_layout);
            # Internal Assessment always targets a fixed 5-pair / 10-question
            # structure regardless of topic tagging (build_fixed_pair_layout)
            # — see that function's docstring for why topic-based grouping
            # is unreliable for IAT papers. Internal Assessment also omits
            # the "Module – N" divider rows (not part of the IAT format)
            # and uses IAT-appropriate note wording.
            if is_internal_assessment:
                vtu_modules = build_fixed_pair_layout(pa.questions, pair_count=5)
            else:
                vtu_modules = build_vtu_paper_layout(pa.questions)
            self._add_vtu_exam_header(
                doc, pa, vtu_modules, show_module_labels=is_semester_exam
            )
            self._add_vtu_question_table(
                doc, pa, vtu_modules, show_module_labels=is_semester_exam
            )
            self._add_co_coverage_table(doc, vtu_modules)
        elif is_case_study:
            self._add_header(doc, pa)
            self._add_title_block(doc, pa)
            self._add_metadata_table(doc, pa)

            if pa.has_instructions:
                self._add_instructions(doc, pa)

            self._add_case_study_questions(doc, pa)
        else:
            self._add_header(doc, pa)
            self._add_title_block(doc, pa)
            self._add_metadata_table(doc, pa)

            if pa.has_instructions:
                self._add_instructions(doc, pa)

            self._add_questions(doc, pa)

        if pa.has_answer_key:
            if is_case_study:
                self._add_answer_key(
                    doc, pa,
                    heading="Evaluation Criteria & Model Approach",
                    entry_label=(
                        "For faculty use only — not to be distributed to "
                        "students. There is rarely a single \"correct\" "
                        "answer for an open-ended case; use this as a "
                        "grading rubric, not a fixed key."
                    ),
                )
            else:
                self._add_answer_key(doc, pa)

        if pa.has_generation_notes:
            self._add_generation_notes(doc, pa)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        raw = buf.read()
        logger.info("WordExporter: produced %d bytes", len(raw))
        return raw

    # ── Document layout ───────────────────────────────────────────────────────

    @staticmethod
    def _configure_page(doc: DocxDocument) -> None:
        """Set A4 page size and 2.5 cm margins."""
        from docx.oxml.ns import nsmap  # noqa: PLC0415
        from docx.oxml import OxmlElement  # noqa: PLC0415

        section = doc.sections[0]
        section.page_height = Inches(11.69)   # A4
        section.page_width = Inches(8.27)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Header / footer ───────────────────────────────────────────────────────

    def _add_header(self, doc: DocxDocument, pa: ParsedAssessment) -> None:
        """Add a branded header with university name and course info."""
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False

        # Clear default paragraph
        for p in header.paragraphs:
            p.clear()

        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_uni = p.add_run(pa.university)
        run_uni.bold = True
        run_uni.font.size = Pt(11)
        run_uni.font.color.rgb = self._COLOUR_PRIMARY

        p.add_run(" · ")

        run_dept = p.add_run(pa.department)
        run_dept.font.size = Pt(10)
        run_dept.font.color.rgb = self._COLOUR_PRIMARY

        # Bottom border
        self._set_paragraph_border(p, bottom_color="1A376C", bottom_size=12)

    @staticmethod
    def _add_footer(doc: DocxDocument, pa: ParsedAssessment) -> None:
        """Add page-numbered footer with assessment title."""
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False

        for p in footer.paragraphs:
            p.clear()

        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run_title = p.add_run(f"{pa.title}  |  ")
        run_title.font.size = Pt(8)

        # Auto page number field
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        p.runs[-1]._r.append(fld)

        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        p.runs[-1]._r.append(instrText)

        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        p.runs[-1]._r.append(fld_end)

        run_of = p.add_run(" of ")
        run_of.font.size = Pt(8)

        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "begin")
        p.runs[-1]._r.append(fld2)

        instrText2 = OxmlElement("w:instrText")
        instrText2.text = "NUMPAGES"
        p.runs[-1]._r.append(instrText2)

        fld2_end = OxmlElement("w:fldChar")
        fld2_end.set(qn("w:fldCharType"), "end")
        p.runs[-1]._r.append(fld2_end)

        run_date = p.add_run(f"  |  Generated: {pa.date_generated}")
        run_date.font.size = Pt(8)

    # ── Content sections ──────────────────────────────────────────────────────

    def _add_title_block(
        self, doc: DocxDocument, pa: ParsedAssessment
    ) -> None:
        """Add the assessment title centred with primary colour."""
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(pa.title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = self._COLOUR_PRIMARY
        doc.add_paragraph()  # spacer

    def _add_metadata_table(
        self, doc: DocxDocument, pa: ParsedAssessment
    ) -> None:
        """Render a two-column metadata info table."""
        rows = [
            ("Course Code", pa.course_code or "—"),
            ("Course Name", pa.course_name or "—"),
            ("Assessment Type", pa.assessment_type),
            ("Semester", pa.semester),
            ("Duration", pa.duration),
            ("Total Marks", pa.total_marks),
            ("Faculty", pa.faculty_name),
            ("Test Date", pa.test_date or "—"),
            ("Generated On", pa.date_generated),
        ]

        table = doc.add_table(rows=len(rows), cols=2)
        table.style = "Table Grid"

        for i, (label, value) in enumerate(rows):
            row = table.rows[i]
            cell_label = row.cells[0]
            cell_value = row.cells[1]

            cell_label.text = label
            cell_value.text = value

            # Style label cell
            run_label = cell_label.paragraphs[0].runs[0]
            run_label.bold = True
            run_label.font.color.rgb = self._COLOUR_PRIMARY
            run_label.font.size = Pt(10)

            cell_value.paragraphs[0].runs[0].font.size = Pt(10)

            # Shade alternate rows
            if i % 2 == 0:
                self._shade_cell(cell_label, "F0F4F8")
                self._shade_cell(cell_value, "F0F4F8")

        # Set column widths
        for row in table.rows:
            row.cells[0].width = Inches(2.0)
            row.cells[1].width = Inches(4.0)

        doc.add_paragraph()

    # ── VTU-style Semester Examination layout ───────────────────────────────

    def _add_vtu_exam_header(
        self,
        doc: DocxDocument,
        pa: ParsedAssessment,
        modules: List[VTUModule],
        show_module_labels: bool = True,
    ) -> None:
        """Render the standard VTU-style OR-pair paper header block.

        Matches the official institutional paper layout: logo / name /
        accreditation lines / IQAC badge, a rule, a centred USN box,
        an ALL-CAPS department heading, a title+metadata table (course,
        code, semester, max marks, batch, duration, IAT date, teaching
        department, RBT legend), and an instruction line — followed by
        (for Semester Examination only) an "answer any N" note, since
        that format spans multiple Modules and needs the extra
        clarification the fixed institutional Internal Assessment
        instruction line doesn't require.

        Args:
            doc: The in-progress python-docx Document.
            pa: The parsed assessment.
            modules: Pre-built VTU module/OR-pair layout (shared with
                :meth:`_add_vtu_question_table` so both reflect the exact
                same structure and the true achievable max marks — only
                ONE alternative per module is ever actually answered, so
                summing every question's marks would double-count).
            show_module_labels: When True (Semester Examination), an
                extra "answer any N, one per MODULE" note is appended.
                When False (Internal Assessment), the fixed
                institutional instruction line is used instead, matching
                the official paper exactly.
        """
        # Date/time stamp, top-left corner — printed before anything else.
        p_stamp = doc.add_paragraph()
        p_stamp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_stamp = p_stamp.add_run(pa.datetime_generated)
        run_stamp.font.size = Pt(8)
        run_stamp.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        p_stamp.paragraph_format.space_after = Pt(2)

        # Institution header: logo | name + affiliation | accreditation
        # lines | IQAC badge. Logo/badge images are optional (see
        # _INSTITUTION_LOGO_PATH / _IQAC_BADGE_PATH) — if the file isn't
        # present, that column is simply left blank rather than breaking
        # the export.
        header_table = doc.add_table(rows=1, cols=4)
        header_table.autofit = False
        # Printable width is 5.87in (8.27in A4 minus 1.2in margins each
        # side, per _configure_page) — these must sum to fit within that,
        # or Word pushes the last column off the page.
        col_widths = [Inches(0.7), Inches(2.6), Inches(1.87), Inches(0.7)]
        for c, w in enumerate(col_widths):
            header_table.columns[c].width = w
            header_table.cell(0, c).width = w
        for row in header_table.rows:
            for cell in row.cells:
                self._set_cell_no_borders(cell)

        # Col 0: institution logo
        cell_logo = header_table.cell(0, 0)
        cell_logo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if _INSTITUTION_LOGO_PATH.exists():
            cell_logo.paragraphs[0].add_run().add_picture(
                str(_INSTITUTION_LOGO_PATH), width=Inches(0.55)
            )

        # Col 1: institution name + affiliation
        cell_name = header_table.cell(0, 1)
        p_uni = cell_name.paragraphs[0]
        p_uni.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_uni = p_uni.add_run(pa.university)
        run_uni.bold = True
        run_uni.font.size = Pt(13)
        run_uni.font.color.rgb = self._COLOUR_PRIMARY
        p_affil = cell_name.add_paragraph()
        p_affil.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_affil = p_affil.add_run(_UNIVERSITY_AFFILIATION)
        run_affil.font.size = Pt(10)
        run_affil.italic = True

        # Col 2: accreditation lines — key terms (VTU, AICTE, NAAC, A+,
        # NBA) rendered in red, matching the official paper's styling.
        cell_accred = header_table.cell(0, 2)
        p_first = cell_accred.paragraphs[0]
        p_first.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for i, line in enumerate(_UNIVERSITY_ACCREDITATION_LINES):
            p = p_first if i == 0 else cell_accred.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            for part in _ACCREDITATION_HIGHLIGHT_RE.split(line):
                if not part:
                    continue
                run = p.add_run(part)
                run.font.size = Pt(8)
                run.font.color.rgb = (
                    RGBColor.from_string(_ACCREDITATION_HIGHLIGHT_COLOR)
                    if part in _ACCREDITATION_HIGHLIGHT_WORDS
                    else RGBColor(0x22, 0x22, 0x22)
                )

        # Col 3: IQAC badge
        cell_badge = header_table.cell(0, 3)
        cell_badge.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if _IQAC_BADGE_PATH.exists():
            cell_badge.paragraphs[0].add_run().add_picture(
                str(_IQAC_BADGE_PATH), width=Inches(0.55)
            )

        # Rule under the header block
        p_rule = doc.add_paragraph()
        p_rule.paragraph_format.space_before = Pt(4)
        p_rule.paragraph_format.space_after = Pt(6)
        self._set_paragraph_border(p_rule, bottom_color="000000", bottom_size=6)

        # USN box — centred (not full-width), matching the official layout.
        usn_table = doc.add_table(rows=1, cols=11)
        usn_table.style = "Table Grid"
        usn_table.autofit = False
        usn_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
        usn_table.cell(0, 0).text = "USN"
        usn_table.cell(0, 0).paragraphs[0].runs[0].bold = True
        usn_table.cell(0, 0).width = Inches(0.6)
        for c in range(1, 11):
            usn_table.cell(0, c).width = Inches(0.35)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

        # Department line — ALL CAPS, bold, left-aligned.
        p_dept = doc.add_paragraph()
        p_dept.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_dept = p_dept.add_run(_department_heading(pa.department).upper())
        run_dept.bold = True
        run_dept.font.size = Pt(11)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # True achievable max marks = sum of ONE alternative per module
        # (the first side of each OR-pair) — NOT a sum of every question,
        # which would double-count both alternatives of every module even
        # though a student only ever answers one of them.
        true_max_marks = sum(
            pair[0].total_marks for m in modules for pair in m.pairs[:1]
        )
        max_marks_display = str(true_max_marks) if true_max_marks else pa.total_marks

        # Title + metadata table: row 0 is a merged, centred title row
        # ("<Assessment Type> (<...>)   AY- <academic_year>"); rows 1-5
        # are the Course/Code/Semester/Max-Marks/Batch/Duration/Date-of-
        # IAT/Teaching-Department/RBT-Levels grid.
        meta = doc.add_table(rows=6, cols=4)
        meta.style = "Table Grid"

        def _set(r: int, c: int, text: str, bold_label: bool = False) -> None:
            cell = meta.cell(r, c)
            cell.text = text
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = bold_label
                cell.paragraphs[0].runs[0].font.size = Pt(10)

        title_bits = [pa.semester, pa.assessment_type]
        title_text = "  ".join(b for b in title_bits if b)
        if pa.academic_year:
            title_text += f"     {pa.academic_year}"
        _set(0, 0, title_text, bold_label=True)
        meta.cell(0, 0).merge(meta.cell(0, 3))
        meta.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        _set(1, 0, "Course:", bold_label=True)
        _set(1, 1, pa.course_name or "—")
        _set(1, 2, "Course Code:", bold_label=True)
        _set(1, 3, pa.course_code or "—")

        _set(2, 0, "Semester:", bold_label=True)
        _set(2, 1, pa.semester or "—")
        _set(2, 2, "Max. Marks:", bold_label=True)
        _set(2, 3, max_marks_display)

        _set(3, 0, "Batch:", bold_label=True)
        _set(3, 1, pa.batch or "—")
        _set(3, 2, "Duration:", bold_label=True)
        _set(3, 3, pa.duration)

        _set(4, 0, "Date of IAT:", bold_label=True)
        _set(4, 1, pa.test_date or "—")
        _set(4, 2, "Teaching Department:", bold_label=True)
        _set(4, 3, pa.teaching_department or "—")

        _set(5, 0, "RBT Levels:", bold_label=True)
        _set(
            5, 1,
            "L1-Remember, L2-Understand, L3-Apply, L4-Analyze, "
            "L5-Evaluate, L6-Create",
        )
        meta.cell(5, 1).merge(meta.cell(5, 3))

        doc.add_paragraph()

        # Instruction line — the fixed institutional wording for Internal
        # Assessment; Semester Examination keeps its dynamic per-module
        # note, since that format needs the extra module-choice
        # clarification a single fixed line can't express.
        p_note = doc.add_paragraph()
        if show_module_labels:
            module_count = len(modules) or 1
            note_text = (
                f"Note: Answer any {module_count} full questions, choosing at "
                f"least ONE question from each MODULE."
            )
            run_note = p_note.add_run(note_text)
            run_note.bold = True
            run_note.font.size = Pt(10)
        else:
            run_note = p_note.add_run(
                "Instruction: Answer the following questions"
            )
            run_note.italic = True
            run_note.bold = True
            run_note.font.size = Pt(10)
        doc.add_paragraph()

    def _add_vtu_question_table(
        self,
        doc: DocxDocument,
        pa: ParsedAssessment,
        modules: List[VTUModule],
        show_module_labels: bool = True,
    ) -> None:
        """Render the Module / OR-pair / sub-part question table.

        Builds one continuous table for the whole paper: a column-header
        row (Q No / Questions / Marks / COs / RBTL — matching the
        official institutional paper's column labels exactly), then for
        each Module (optionally) a full-width "Module – N" divider row
        followed by its question groups, with a full-width "OR" divider
        row between the two alternatives in a pair. Each lettered
        sub-part is its own full row labelled "N. a", "N. b", … (no
        separate letter column, no vertical cell merging) — matching the
        official paper's row format exactly; a single un-split question
        is labelled just "N.".

        Args:
            doc: The in-progress python-docx Document.
            pa: The parsed assessment.
            modules: Pre-built VTU module/OR-pair layout (see
                :func:`build_vtu_paper_layout`), shared with
                :meth:`_add_vtu_exam_header` for consistency.
            show_module_labels: When True (Semester Examination), print a
                "Module – N: label" divider row before each module's
                pair(s). When False (Internal Assessment), omit these
                rows so the paper flows as a continuous Q1 OR Q2, Q3 OR
                Q4, … sequence, matching the institutional IAT format.
        """
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["Q No", "Questions", "Marks", "COs", "RBTL"]
        for c, text in enumerate(headers):
            cell = table.cell(0, c)
            cell.text = text
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
            self._shade_cell(cell, "F0F4F8")

        col_widths = [Inches(w) for w in (0.55, 3.9, 0.6, 0.6, 0.6)]

        def _full_width_row(text: str, bold: bool = True) -> None:
            row_cells = table.add_row().cells
            row_cells[0].merge(row_cells[-1])
            row_cells[0].text = text
            p = row_cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].bold = bold
            self._shade_cell(row_cells[0], "F0F4F8")

        def _question_group_rows(group: VTUQuestionGroup) -> None:
            multi = len(group.subparts) > 1
            for sp in group.subparts:
                row_cells = table.add_row().cells
                q_label = f"{group.q_number}. {sp.letter}" if multi else f"{group.q_number}."
                row_cells[0].text = q_label
                row_cells[0].paragraphs[0].runs[0].bold = True
                # question text
                row_cells[1].text = sp.question.text
                for run in row_cells[1].paragraphs[0].runs:
                    run.font.size = Pt(10)
                # marks / CO / RBTL
                row_cells[2].text = sp.question.marks
                row_cells[3].text = sp.question.co_mapping
                row_cells[4].text = self._bloom_short_code(sp.question.bloom_level)
                for idx in (0, 2, 3, 4):
                    if row_cells[idx].paragraphs[0].runs:
                        row_cells[idx].paragraphs[0].runs[0].font.size = Pt(10)
                    row_cells[idx].paragraphs[0].alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                    )

        for module in modules:
            if show_module_labels:
                _full_width_row(f"Module – {module.module_number}: {module.label}")
            for pair_idx, pair in enumerate(module.pairs):
                if pair_idx > 0:
                    _full_width_row("OR")
                _mirror_pair_co_bloom(pair)
                _question_group_rows(pair[0])
                if len(pair) > 1:
                    _full_width_row("OR")
                    _question_group_rows(pair[1])

        for row in table.rows:
            for c, cell in enumerate(row.cells):
                if c < len(col_widths):
                    cell.width = col_widths[c]

        doc.add_paragraph()

    def _add_co_coverage_table(
        self, doc: DocxDocument, modules: List["VTUModule"]
    ) -> None:
        """Render the CO coverage summary table right after the questions.

        A two-row table: "Course Outcomes" header row listing every CO
        that appears in the paper, and a "Percentage" row with each CO's
        marks-weighted share of the paper's total (achievable) marks —
        see :func:`_compute_co_coverage` for how that's calculated.
        Renders nothing when there's no CO data to show.

        Args:
            doc: The in-progress python-docx Document.
            modules: Pre-built VTU module/OR-pair layout.
        """
        coverage = _compute_co_coverage(modules)
        if not coverage:
            return

        table = doc.add_table(rows=2, cols=len(coverage) + 1)
        table.style = "Table Grid"

        table.cell(0, 0).text = "Course Outcomes"
        table.cell(0, 0).paragraphs[0].runs[0].bold = True
        table.cell(1, 0).text = "Percentage"
        table.cell(1, 0).paragraphs[0].runs[0].bold = True

        for i, (co, pct) in enumerate(coverage, start=1):
            cell_co = table.cell(0, i)
            cell_co.text = co
            cell_co.paragraphs[0].runs[0].bold = True
            cell_co.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            cell_pct = table.cell(1, i)
            cell_pct.text = f"{pct}%"
            cell_pct.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

    @staticmethod
    def _bloom_short_code(bloom_level: str) -> str:
        """Map a full Bloom's level name to its short "L1"-"L6" code."""
        order = [
            "Remember", "Understand", "Apply", "Analyze", "Evaluate", "Create",
        ]
        try:
            return f"L{order.index(bloom_level) + 1}"
        except ValueError:
            return bloom_level[:2] if bloom_level else "—"

    def _add_instructions(
        self, doc: DocxDocument, pa: ParsedAssessment
    ) -> None:
        """Add a shaded instructions box."""
        h = doc.add_heading("Instructions", level=2)
        h.runs[0].font.color.rgb = self._COLOUR_PRIMARY

        p = doc.add_paragraph(pa.instructions)
        p.style = "Normal"
        p.paragraph_format.left_indent = Inches(0.3)
        doc.add_paragraph()

    def _add_questions(
        self, doc: DocxDocument, pa: ParsedAssessment
    ) -> None:
        """Add numbered questions with metadata sub-tables."""
        h = doc.add_heading("Questions", level=1)
        h.runs[0].font.color.rgb = self._COLOUR_PRIMARY
        doc.add_paragraph()

        for q in pa.questions:
            # Question heading
            p_q = doc.add_paragraph()
            run_num = p_q.add_run(f"{q.number}.  ")
            run_num.bold = True
            run_num.font.size = Pt(11)
            run_num.font.color.rgb = self._COLOUR_PRIMARY

            run_text = p_q.add_run(q.text)
            run_text.font.size = Pt(11)

            # MCQ options — one per line with a selection circle
            for i, opt in enumerate(q.options):
                p_opt = doc.add_paragraph(f"○  {chr(65 + i)}.  {opt}")
                p_opt.paragraph_format.left_indent = Inches(0.4)
                for run in p_opt.runs:
                    run.font.size = Pt(10.5)

            # Metadata bar (inline)
            meta_parts = [
                f"[{q.question_type}]",
                f"Bloom: {q.bloom_level}",
                f"Difficulty: {q.difficulty}",
                f"CO: {q.co_mapping}",
                f"Marks: {q.marks}",
            ]
            p_meta = doc.add_paragraph("    ".join(meta_parts))
            p_meta.paragraph_format.left_indent = Inches(0.25)
            for run in p_meta.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

            if q.notes:
                p_note = doc.add_paragraph(f"Note: {q.notes}")
                p_note.paragraph_format.left_indent = Inches(0.25)
                for run in p_note.runs:
                    run.italic = True
                    run.font.size = Pt(9)

            doc.add_paragraph()

    def _add_case_study_questions(
        self, doc: DocxDocument, pa: ParsedAssessment
    ) -> None:
        """Render Consultancy Case Study questions in a branded layout.

        Each case gets its own "Case N" heading, followed by clearly
        labelled "Case Background" and "Your Task" sections (rather than
        the generic single question_text block), matching a real
        consulting case-study handout rather than an exam question list.
        """
        h = doc.add_heading("Cases", level=1)
        h.runs[0].font.color.rgb = self._COLOUR_PRIMARY
        doc.add_paragraph()

        for i, q in enumerate(pa.questions):
            if i > 0:
                doc.add_page_break()

            p_case = doc.add_paragraph()
            run_case = p_case.add_run(f"Case {q.number}")
            run_case.bold = True
            run_case.font.size = Pt(15)
            run_case.font.color.rgb = self._COLOUR_PRIMARY
            doc.add_paragraph()

            if q.case_background:
                h_bg = doc.add_heading("Case Background", level=2)
                h_bg.runs[0].font.color.rgb = self._COLOUR_PRIMARY
                p_bg = doc.add_paragraph(q.case_background)
                p_bg.paragraph_format.space_after = Pt(10)
                doc.add_paragraph()

            h_task = doc.add_heading("Your Task", level=2)
            h_task.runs[0].font.color.rgb = self._COLOUR_PRIMARY
            p_task = doc.add_paragraph(q.text)
            p_task.paragraph_format.space_after = Pt(10)
            doc.add_paragraph()

            meta_parts = [
                f"[{q.question_type}]",
                f"Bloom: {q.bloom_level}",
                f"Difficulty: {q.difficulty}",
                f"CO: {q.co_mapping}",
                f"Marks: {q.marks}",
            ]
            p_meta = doc.add_paragraph("    ".join(meta_parts))
            p_meta.paragraph_format.left_indent = Inches(0.25)
            for run in p_meta.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x55, 0x66, 0x77)

            if q.notes:
                p_note = doc.add_paragraph(f"Note: {q.notes}")
                p_note.paragraph_format.left_indent = Inches(0.25)
                for run in p_note.runs:
                    run.italic = True
                    run.font.size = Pt(9)

            doc.add_paragraph()

    def _add_answer_key(
        self,
        doc: DocxDocument,
        pa: ParsedAssessment,
        heading: str = "Answer Key",
        entry_label: str = "For faculty use only — not to be distributed to students.",
    ) -> None:
        """Add answer-key section after a page break.

        Args:
            doc: The in-progress python-docx Document.
            pa: The parsed assessment.
            heading: Section heading text — overridden to "Evaluation
                Criteria & Model Approach" for Consultancy Case Studies,
                since there's rarely a single fixed "correct answer" for
                an open-ended business case the way there is for a
                typical exam question.
            entry_label: The faculty-only notice line under the heading.
        """
        doc.add_page_break()

        h = doc.add_heading(heading, level=1)
        h.runs[0].font.color.rgb = self._COLOUR_PRIMARY

        notice = doc.add_paragraph(entry_label)
        for run in notice.runs:
            run.italic = True
            run.font.color.rgb = RGBColor(0xAA, 0x44, 0x00)
        doc.add_paragraph()

        for q in pa.questions:
            p_q = doc.add_paragraph()
            run_num = p_q.add_run(f"{q.number}. {q.question_id}  ")
            run_num.bold = True
            run_num.font.size = Pt(11)
            run_num.font.color.rgb = self._COLOUR_PRIMARY

            answer_text = q.answer_key if q.answer_key else "Not provided."
            p_ans = doc.add_paragraph(answer_text)
            p_ans.paragraph_format.left_indent = Inches(0.3)
            doc.add_paragraph()

    def _add_generation_notes(
        self, doc: DocxDocument, pa: ParsedAssessment
    ) -> None:
        """Add agent generation notes at the end."""
        h = doc.add_heading("Agent Notes", level=2)
        h.runs[0].font.color.rgb = self._COLOUR_PRIMARY
        doc.add_paragraph(pa.generation_notes)

    # ── XML helpers ───────────────────────────────────────────────────────────

    @staticmethod
    def _set_cell_no_borders(cell) -> None:
        """Remove all borders from a table cell (for borderless layout tables)."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge in ("top", "left", "bottom", "right"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
            borders.append(el)
        tcPr.append(borders)

    @staticmethod
    def _shade_cell(cell, fill_hex: str) -> None:
        """Apply a solid background shade to a table cell."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill_hex)
        tcPr.append(shd)

    @staticmethod
    def _set_paragraph_border(
        paragraph,
        bottom_color: str = "000000",
        bottom_size: int = 6,
    ) -> None:
        """Add a bottom border to a paragraph via direct XML manipulation."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), str(bottom_size))
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), bottom_color)
        pBdr.append(bottom)
        pPr.append(pBdr)


# ===========================================================================
# PDF exporter
# ===========================================================================


class PDFExporter:
    """Export an :class:`~models.Assessment` to a professional PDF byte stream.

    Uses ``ReportLab`` with the parsed intermediate representation.
    Documents are built from typed :class:`ParsedAssessment` objects — never
    from raw LLM output strings — to avoid ``Paragraph()`` XML-parsing issues.

    Layout features:
    - A4 page with 2.5 cm margins
    - University header (coloured rule + name) and footer (title + page numbers)
    - Metadata grid table
    - Numbered questions with metadata tag line
    - Answer-key section separated by a page break
    """

    # Colour palette (RGB 0–1 floats for ReportLab)
    _NAVY = colors.Color(0x1A / 255, 0x37 / 255, 0x6C / 255)
    _GOLD = colors.Color(0xC8 / 255, 0x9B / 255, 0x20 / 255)
    _LIGHT_BG = colors.Color(0xF0 / 255, 0xF4 / 255, 0xF8 / 255)
    _GREY_TEXT = colors.Color(0x55 / 255, 0x66 / 255, 0x77 / 255)
    _ANSWER_NOTICE_COLOR = colors.Color(0.6, 0.2, 0.0)

    def export(self, assessment: Assessment) -> bytes:
        """Render the assessment as a PDF byte stream.

        Args:
            assessment: The typed Assessment to export.

        Returns:
            bytes: Raw PDF bytes suitable for ``st.download_button``.

        Raises:
            ValueError: If *assessment* is None.
        """
        if assessment is None:
            raise ValueError("assessment must not be None")

        logger.info(
            "PDFExporter: exporting '%s'", assessment.metadata.title
        )

        pa = AssessmentParser.parse(assessment)
        styles = self._build_styles()

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            rightMargin=2.5 * cm,
            leftMargin=2.5 * cm,
            topMargin=2.5 * cm,
            bottomMargin=2.5 * cm,
            title=pa.title,
            author=pa.faculty_name,
            subject=f"{pa.course_code} {pa.assessment_type}",
        )

        story = self._build_story(pa, styles)

        doc.build(
            story,
            onFirstPage=self._make_page_callback(pa),
            onLaterPages=self._make_page_callback(pa),
        )

        buf.seek(0)
        raw = buf.read()
        logger.info("PDFExporter: produced %d bytes", len(raw))
        return raw

    # ── Style sheet ───────────────────────────────────────────────────────────

    def _build_styles(self) -> dict:
        """Build and return a dict of named ParagraphStyle objects."""
        base = getSampleStyleSheet()

        return {
            "title": ParagraphStyle(
                "EduTitle",
                parent=base["Heading1"],
                fontSize=18,
                textColor=self._NAVY,
                alignment=TA_CENTER,
                spaceAfter=6,
                fontName="Helvetica-Bold",
            ),
            "university": ParagraphStyle(
                "EduUniversity",
                parent=base["Normal"],
                fontSize=11,
                textColor=self._NAVY,
                alignment=TA_CENTER,
                fontName="Helvetica-Bold",
                spaceAfter=2,
            ),
            "vtu_affiliation": ParagraphStyle(
                "EduVTUAffil",
                parent=base["Normal"],
                fontSize=9.5,
                alignment=TA_CENTER,
                fontName="Helvetica-Oblique",
                spaceAfter=6,
            ),
            "vtu_exam_title": ParagraphStyle(
                "EduVTUExamTitle",
                parent=base["Normal"],
                fontSize=11,
                alignment=TA_CENTER,
                fontName="Helvetica-Bold",
                spaceAfter=10,
            ),
            "vtu_note": ParagraphStyle(
                "EduVTUNote",
                parent=base["Normal"],
                fontSize=9.5,
                fontName="Helvetica-Bold",
                spaceAfter=10,
            ),
            "vtu_cell_text": ParagraphStyle(
                "EduVTUCellText",
                parent=base["Normal"],
                fontSize=9,
                leading=12,
                fontName="Helvetica",
            ),
            "vtu_cell_center": ParagraphStyle(
                "EduVTUCellCenter",
                parent=base["Normal"],
                fontSize=9,
                leading=12,
                fontName="Helvetica",
                alignment=TA_CENTER,
            ),
            "vtu_module_header": ParagraphStyle(
                "EduVTUModuleHeader",
                parent=base["Normal"],
                fontSize=10,
                fontName="Helvetica-Bold",
                alignment=TA_CENTER,
            ),
            "department": ParagraphStyle(
                "EduDept",
                parent=base["Normal"],
                fontSize=9,
                textColor=self._NAVY,
                alignment=TA_CENTER,
                fontName="Helvetica",
                spaceAfter=4,
            ),
            "section_heading": ParagraphStyle(
                "EduSection",
                parent=base["Heading2"],
                fontSize=13,
                textColor=self._NAVY,
                fontName="Helvetica-Bold",
                spaceBefore=14,
                spaceAfter=6,
            ),
            "question_heading": ParagraphStyle(
                "EduQ",
                parent=base["Normal"],
                fontSize=11,
                leading=15,
                fontName="Helvetica-Bold",
                textColor=colors.black,
                spaceBefore=10,
                spaceAfter=2,
            ),
            "question_text": ParagraphStyle(
                "EduQText",
                parent=base["Normal"],
                fontSize=11,
                leading=15,
                fontName="Helvetica",
                textColor=colors.black,
                spaceAfter=2,
            ),
            "option_line": ParagraphStyle(
                "EduOption",
                parent=base["Normal"],
                fontSize=10.5,
                leading=14,
                fontName="Helvetica",
                textColor=colors.black,
                leftIndent=18,
                spaceAfter=1,
            ),
            "meta_tag": ParagraphStyle(
                "EduMeta",
                parent=base["Normal"],
                fontSize=8.5,
                textColor=self._GREY_TEXT,
                fontName="Helvetica",
                spaceAfter=4,
                leftIndent=12,
            ),
            "meta_cell_label": ParagraphStyle(
                "EduMetaCellLabel",
                parent=base["Normal"],
                fontSize=9.5,
                fontName="Helvetica-Bold",
                textColor=colors.black,
                leading=12,
            ),
            "meta_cell_value": ParagraphStyle(
                "EduMetaCellValue",
                parent=base["Normal"],
                fontSize=9.5,
                fontName="Helvetica",
                textColor=colors.black,
                leading=12,
            ),
            "note_tag": ParagraphStyle(
                "EduNote",
                parent=base["Normal"],
                fontSize=9,
                textColor=self._GREY_TEXT,
                fontName="Helvetica-Oblique",
                spaceAfter=4,
                leftIndent=12,
            ),
            "answer_notice": ParagraphStyle(
                "EduAnsNotice",
                parent=base["Normal"],
                fontSize=9,
                textColor=self._ANSWER_NOTICE_COLOR,
                fontName="Helvetica-Oblique",
                alignment=TA_CENTER,
                spaceBefore=4,
                spaceAfter=10,
            ),
            "answer_id": ParagraphStyle(
                "EduAnsId",
                parent=base["Normal"],
                fontSize=11,
                fontName="Helvetica-Bold",
                textColor=self._NAVY,
                spaceBefore=8,
                spaceAfter=2,
            ),
            "answer_body": ParagraphStyle(
                "EduAnsBody",
                parent=base["Normal"],
                fontSize=10,
                fontName="Helvetica",
                textColor=colors.black,
                leading=14,
                leftIndent=16,
                spaceAfter=4,
                alignment=TA_JUSTIFY,
            ),
            "footer": ParagraphStyle(
                "EduFooter",
                parent=base["Normal"],
                fontSize=8,
                textColor=self._GREY_TEXT,
                alignment=TA_CENTER,
                fontName="Helvetica",
            ),
            "agent_notes": ParagraphStyle(
                "EduAgentNotes",
                parent=base["Normal"],
                fontSize=10,
                fontName="Helvetica-Oblique",
                textColor=self._GREY_TEXT,
                leading=14,
                spaceAfter=4,
            ),
            "normal": base["Normal"],
        }

    # ── Story builder ─────────────────────────────────────────────────────────

    def _build_story(
        self, pa: ParsedAssessment, styles: dict
    ) -> list:
        """Assemble the ReportLab story (flowable list) from ParsedAssessment.

        All text is sanitised through :meth:`_safe_para` so that stray
        HTML-like characters (``<``, ``>``, ``&``) never cause a Paragraph
        parse error.

        Args:
            pa: The parsed assessment.
            styles: Dict of named ParagraphStyle instances.

        Returns:
            list: List of ReportLab Flowable objects.
        """
        story = []
        is_semester_exam = pa.assessment_type == AssessmentType.SEMESTER_EXAM.value
        is_internal_assessment = pa.assessment_type == AssessmentType.INTERNAL.value
        is_case_study = (
            pa.assessment_type == AssessmentType.CONSULTANCY_CASE.value
        )

        if is_semester_exam or is_internal_assessment:
            if is_internal_assessment:
                vtu_modules = build_fixed_pair_layout(pa.questions, pair_count=5)
            else:
                vtu_modules = build_vtu_paper_layout(pa.questions)
            story.extend(
                self._build_vtu_header(
                    pa, styles, vtu_modules, show_module_labels=is_semester_exam
                )
            )
            story.extend(
                self._build_vtu_question_table(
                    pa, styles, vtu_modules, show_module_labels=is_semester_exam
                )
            )
            story.extend(self._build_co_coverage_table(vtu_modules, styles))
        else:
            # ── Title block ──────────────────────────────────────────────
            story.append(Paragraph(self._esc(pa.university), styles["university"]))
            story.append(Paragraph(self._esc(pa.department), styles["department"]))
            story.append(HRFlowable(width="100%", thickness=1.5, color=self._NAVY))
            story.append(Spacer(1, 8))
            story.append(Paragraph(self._esc(pa.title), styles["title"]))
            story.append(Spacer(1, 6))

            # ── Metadata table ───────────────────────────────────────────
            meta_data = [
                ["Course Code", pa.course_code or "—",
                 "Assessment Type", pa.assessment_type],
                ["Course Name", pa.course_name or "—",
                 "Semester", pa.semester],
                ["Duration", pa.duration,
                 "Total Marks", pa.total_marks],
                ["Faculty", pa.faculty_name,
                 "Test Date", pa.test_date or "—"],
                ["Generated On", pa.date_generated,
                 "", ""],
            ]
            meta_table = Table(
                meta_data,
                colWidths=[3.5 * cm, 5.5 * cm, 3.5 * cm, 5.5 * cm],
            )
            meta_table.setStyle(
                TableStyle([
                    ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                    ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
                    ("TEXTCOLOR", (0, 0), (0, -1), self._NAVY),
                    ("TEXTCOLOR", (2, 0), (2, -1), self._NAVY),
                    ("BACKGROUND", (0, 0), (-1, -1), self._LIGHT_BG),
                    ("ROWBACKGROUNDS", (0, 0), (-1, -1),
                     [self._LIGHT_BG, colors.white]),
                    ("BOX", (0, 0), (-1, -1), 0.5, self._NAVY),
                    ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ])
            )
            story.append(meta_table)
            story.append(Spacer(1, 10))

            # ── Instructions ─────────────────────────────────────────────
            if pa.has_instructions:
                story.append(
                    Paragraph("Instructions", styles["section_heading"])
                )
                story.append(
                    Paragraph(self._esc(pa.instructions), styles["normal"])
                )
                story.append(Spacer(1, 6))

            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))

            # ── Questions ────────────────────────────────────────────────
            if is_case_study:
                story.append(Paragraph("Cases", styles["section_heading"]))
                for i, q in enumerate(pa.questions):
                    if i > 0:
                        story.append(PageBreak())
                    story.append(Paragraph(f"Case {q.number}", styles["title"]))
                    story.append(Spacer(1, 6))
                    if q.case_background:
                        story.append(
                            Paragraph("Case Background", styles["section_heading"])
                        )
                        story.append(
                            Paragraph(self._esc(q.case_background), styles["normal"])
                        )
                        story.append(Spacer(1, 8))
                    story.append(
                        Paragraph("Your Task", styles["section_heading"])
                    )
                    story.append(
                        Paragraph(self._esc(q.text), styles["normal"])
                    )
                    story.append(Spacer(1, 8))
                    meta_line = (
                        f"[{q.question_type}]  •  Bloom: {q.bloom_level}  •  "
                        f"Difficulty: {q.difficulty}  •  CO: {q.co_mapping}  •  "
                        f"Marks: {q.marks}"
                    )
                    story.append(
                        Paragraph(self._esc(meta_line), styles["meta_tag"])
                    )
                    if q.notes:
                        story.append(
                            Paragraph(
                                f"Note: {self._esc(q.notes)}",
                                styles["note_tag"],
                            )
                        )
            else:
                story.append(Paragraph("Questions", styles["section_heading"]))

                for q in pa.questions:
                    meta_line = (
                        f"[{q.question_type}]  •  Bloom: {q.bloom_level}  •  "
                        f"Difficulty: {q.difficulty}  •  CO: {q.co_mapping}  •  "
                        f"Marks: {q.marks}"
                    )
                    block = [
                        Paragraph(
                            f"{q.number}.  {self._esc(q.text)}",
                            styles["question_text"],
                        ),
                    ]
                    for i, opt in enumerate(q.options):
                        block.append(
                            Paragraph(
                                f"○&nbsp;&nbsp;{chr(65 + i)}.&nbsp;&nbsp;"
                                f"{self._esc(opt)}",
                                styles["option_line"],
                            )
                        )
                    block.append(Paragraph(self._esc(meta_line), styles["meta_tag"]))
                    if q.notes:
                        block.append(
                            Paragraph(f"Note: {self._esc(q.notes)}", styles["note_tag"])
                        )
                    story.append(KeepTogether(block))

        # ── Answer key ───────────────────────────────────────────────────────
        if pa.has_answer_key:
            story.append(Spacer(1, 6))
            story.append(HRFlowable(width="100%", thickness=1, color=self._NAVY))
            if is_case_study:
                story.append(
                    Paragraph(
                        "Evaluation Criteria & Model Approach",
                        styles["section_heading"],
                    )
                )
                story.append(
                    Paragraph(
                        "For faculty use only — not to be distributed to "
                        "students. There is rarely a single \"correct\" "
                        "answer for an open-ended case; use this as a "
                        "grading rubric, not a fixed key.",
                        styles["answer_notice"],
                    )
                )
            else:
                story.append(Paragraph("Answer Key", styles["section_heading"]))
                story.append(
                    Paragraph(
                        "For faculty use only — not to be distributed to students.",
                        styles["answer_notice"],
                    )
                )

            for q in pa.questions:
                answer_text = q.answer_key if q.answer_key else "Not provided."
                ak_block = [
                    Paragraph(
                        f"{q.number}. {self._esc(q.question_id)}",
                        styles["answer_id"],
                    ),
                    Paragraph(
                        self._esc(answer_text), styles["answer_body"]
                    ),
                ]
                story.append(KeepTogether(ak_block))

        # ── Generation notes ─────────────────────────────────────────────────
        if pa.has_generation_notes:
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.lightgrey))
            story.append(Paragraph("Agent Notes", styles["section_heading"]))
            story.append(
                Paragraph(self._esc(pa.generation_notes), styles["agent_notes"])
            )

        return story

    # ── VTU-style Semester Examination layout ───────────────────────────────

    def _build_vtu_header(
        self,
        pa: ParsedAssessment,
        styles: dict,
        modules: list,
        show_module_labels: bool = True,
    ) -> list:
        """Build the USN box, institution header, and course meta table.

        Args:
            pa: The parsed assessment.
            styles: Named ParagraphStyle dict.
            modules: Pre-built VTU module/OR-pair layout, shared with
                :meth:`_build_vtu_question_table` — used here to compute
                the true achievable max marks (one alternative per
                module, not a sum of every printed question) and the
                accurate module count for the note line.
            show_module_labels: When True (Semester Examination), the
                note refers to "each MODULE". When False (Internal
                Assessment), the note refers to "each OR pair" instead.

        Returns:
            list: Flowables for the exam paper header block.
        """
        flow: list = []

        # Date/time stamp, top-left corner — printed before anything else.
        flow.append(Paragraph(self._esc(pa.datetime_generated), styles["meta_tag"]))
        flow.append(Spacer(1, 2))

        # Institution header: logo | name + affiliation | accreditation
        # lines (key terms in red) | IQAC badge. Missing image files
        # degrade gracefully to a blank cell.
        name_block = [
            Paragraph(self._esc(pa.university), styles["university"]),
            Paragraph(
                self._esc(_UNIVERSITY_AFFILIATION), styles["vtu_affiliation"]
            ),
        ]
        accred_block = [
            Paragraph(self._accreditation_html(line), styles["meta_tag"])
            for line in _UNIVERSITY_ACCREDITATION_LINES
        ]
        logo_cell = (
            Image(str(_INSTITUTION_LOGO_PATH), width=0.55 * 72, height=0.55 * 72)
            if _INSTITUTION_LOGO_PATH.exists()
            else ""
        )
        badge_cell = (
            Image(str(_IQAC_BADGE_PATH), width=0.55 * 72, height=0.55 * 72)
            if _IQAC_BADGE_PATH.exists()
            else ""
        )
        # Printable width is 16cm (A4 21cm minus 2.5cm margins each side,
        # per SimpleDocTemplate below) — these must sum to fit within that.
        header_table = Table(
            [[logo_cell, name_block, accred_block, badge_cell]],
            colWidths=[2.0 * cm, 7.0 * cm, 5.0 * cm, 2.0 * cm],
        )
        header_table.setStyle(
            TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (0, 0), "CENTER"),
                ("ALIGN", (1, 0), (1, 0), "LEFT"),
                ("ALIGN", (2, 0), (2, 0), "LEFT"),
                ("ALIGN", (3, 0), (3, 0), "CENTER"),
                ("LEFTPADDING", (0, 0), (-1, -1), 2),
                ("RIGHTPADDING", (0, 0), (-1, -1), 2),
            ])
        )
        flow.append(header_table)
        flow.append(Spacer(1, 4))
        flow.append(HRFlowable(width="100%", thickness=1.2, color=colors.black))
        flow.append(Spacer(1, 6))

        # USN entry box — one wide label cell + 10 individual digit boxes,
        # right-aligned to match the official paper's positioning.
        usn_data = [["USN"] + [""] * 10]
        usn_table = Table(
            usn_data, colWidths=[1.6 * cm] + [0.9 * cm] * 10, rowHeights=[0.7 * cm],
            hAlign="RIGHT",
        )
        usn_table.setStyle(
            TableStyle([
                ("BOX", (0, 0), (-1, -1), 0.75, colors.black),
                ("INNERGRID", (0, 0), (-1, -1), 0.75, colors.black),
                ("FONTNAME", (0, 0), (0, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (0, 0), 5),
            ])
        )
        flow.append(usn_table)
        flow.append(Spacer(1, 10))

        # Department line — ALL CAPS, bold, left-aligned.
        flow.append(
            Paragraph(
                self._esc(_department_heading(pa.department).upper()),
                styles["vtu_exam_title"],
            )
        )
        flow.append(Spacer(1, 8))

        # True achievable max marks = sum of ONE alternative per module
        # (the first side of each OR-pair) — NOT a sum of every question,
        # which would double-count both alternatives of every module even
        # though a student only ever answers one of them.
        true_max_marks = sum(
            pair[0].total_marks for m in modules for pair in m.pairs[:1]
        )
        max_marks_display = str(true_max_marks) if true_max_marks else pa.total_marks

        title_bits = [pa.semester, pa.assessment_type]
        title_text = "  ".join(b for b in title_bits if b)
        if pa.academic_year:
            title_text += f"     {pa.academic_year}"

        # Title + metadata table: row 0 is a merged, centred title row;
        # rows 1-5 are the Course/Code/Semester/Max-Marks/Batch/Duration/
        # Date-of-IAT/Teaching-Department/RBT-Levels grid.
        def _lbl(text: str):
            return Paragraph(self._esc(text), styles["meta_cell_label"])

        def _val(text: str):
            return Paragraph(self._esc(text), styles["meta_cell_value"])

        meta_data = [
            [Paragraph(f"<b>{self._esc(title_text)}</b>", styles["meta_cell_label"]), "", "", ""],
            [_lbl("Course:"), _val(pa.course_name or "—"), _lbl("Course Code:"), _val(pa.course_code or "—")],
            [_lbl("Semester:"), _val(pa.semester or "—"), _lbl("Max. Marks:"), _val(max_marks_display)],
            [_lbl("Batch:"), _val(pa.batch or "—"), _lbl("Duration:"), _val(pa.duration)],
            [_lbl("Date of IAT:"), _val(pa.test_date or "—"), _lbl("Teaching Department:"), _val(pa.teaching_department or "—")],
            [
                _lbl("RBT Levels:"),
                _val(
                    "L1-Remember, L2-Understand, L3-Apply, L4-Analyze, "
                    "L5-Evaluate, L6-Create"
                ),
                "", "",
            ],
        ]
        meta_table = Table(
            meta_data, colWidths=[3.0 * cm, 5.6 * cm, 3.6 * cm, 3.8 * cm]
        )
        meta_table.setStyle(
            TableStyle([
                ("BOX", (0, 0), (-1, -1), 0.75, colors.black),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("SPAN", (0, 0), (3, 0)),
                ("SPAN", (1, 5), (3, 5)),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ])
        )
        flow.append(meta_table)
        flow.append(Spacer(1, 10))

        module_count = len(modules) or 1
        if show_module_labels:
            note_text = (
                f"Note: Answer any {module_count} full questions, choosing "
                f"at least ONE question from each MODULE."
            )
            flow.append(Paragraph(note_text, styles["vtu_note"]))
        else:
            flow.append(
                Paragraph(
                    "<i>Instruction: Answer the following questions</i>",
                    styles["vtu_note"],
                )
            )
        flow.append(Spacer(1, 6))
        return flow

    def _build_vtu_question_table(
        self,
        pa: ParsedAssessment,
        styles: dict,
        modules: list,
        show_module_labels: bool = True,
    ) -> list:
        """Build the Module / OR-pair / sub-part question table.

        Uses ReportLab ``SPAN`` commands to merge the full-width Module
        and "OR" divider rows, and to vertically merge the Q-number cell
        across a question's lettered sub-part rows.

        Args:
            pa: The parsed assessment.
            styles: Named ParagraphStyle dict.
            modules: Pre-built VTU module/OR-pair layout, shared with
                :meth:`_build_vtu_header`.
            show_module_labels: When True (Semester Examination), print
                a "Module – N: label" divider row before each module's
                pair(s). When False (Internal Assessment), omit these
                rows, matching the institutional IAT format.

        Returns:
            list: A single-element list containing the assembled Table
            flowable (wrapped in a list for consistency with other
            flowable-returning builders).
        """
        data: list = [["Q No", "Questions", "Marks", "COs", "RBTL"]]
        style_cmds = [
            ("BACKGROUND", (0, 0), (-1, 0), self._LIGHT_BG),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ]
        row = 1

        def _full_width_row(text: str) -> None:
            nonlocal row
            data.append([text, "", "", "", ""])
            style_cmds.append(("SPAN", (0, row), (-1, row)))
            style_cmds.append(("BACKGROUND", (0, row), (-1, row), self._LIGHT_BG))
            style_cmds.append(("FONTNAME", (0, row), (-1, row), "Helvetica-Bold"))
            style_cmds.append(("ALIGN", (0, row), (-1, row), "CENTER"))
            row += 1

        def _group_rows(group: VTUQuestionGroup) -> None:
            nonlocal row
            multi = len(group.subparts) > 1
            for sp in group.subparts:
                q_label = f"{group.q_number}. {sp.letter}" if multi else f"{group.q_number}."
                data.append([
                    q_label,
                    Paragraph(self._esc(sp.question.text), styles["vtu_cell_text"]),
                    sp.question.marks,
                    sp.question.co_mapping,
                    self._bloom_short_code(sp.question.bloom_level),
                ])
                style_cmds.append(("ALIGN", (0, row), (0, row), "CENTER"))
                style_cmds.append(("VALIGN", (0, row), (0, row), "MIDDLE"))
                style_cmds.append(("FONTNAME", (0, row), (0, row), "Helvetica-Bold"))
                for col in (2, 3, 4):
                    style_cmds.append(("ALIGN", (col, row), (col, row), "CENTER"))
                    style_cmds.append(("VALIGN", (col, row), (col, row), "MIDDLE"))
                row += 1

        for module in modules:
            if show_module_labels:
                _full_width_row(f"Module – {module.module_number}: {module.label}")
            for pair in module.pairs:
                _mirror_pair_co_bloom(pair)
                _group_rows(pair[0])
                if len(pair) > 1:
                    _full_width_row("OR")
                    _group_rows(pair[1])

        col_widths = [
            1.4 * cm, 10.6 * cm, 1.5 * cm, 1.5 * cm, 1.5 * cm,
        ]
        table = Table(data, colWidths=col_widths, repeatRows=1)
        style_cmds.extend([
            ("BOX", (0, 0), (-1, -1), 0.75, colors.black),
            ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.black),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ])
        table.setStyle(TableStyle(style_cmds))
        return [table]

    def _build_co_coverage_table(
        self, modules: List["VTUModule"], styles: dict
    ) -> list:
        """Build the CO coverage summary table right after the questions.

        A two-row table: "Course Outcomes" header row listing every CO
        that appears in the paper, and a "Percentage" row with each CO's
        marks-weighted share of the paper's total (achievable) marks —
        see :func:`_compute_co_coverage` for how that's calculated.
        Returns an empty list when there's no CO data to show.

        Args:
            modules: Pre-built VTU module/OR-pair layout.
            styles: Named ParagraphStyle dict.

        Returns:
            list: Flowables for the CO coverage table (empty if no data).
        """
        coverage = _compute_co_coverage(modules)
        if not coverage:
            return []

        header_row = ["Course Outcomes"] + [co for co, _ in coverage]
        pct_row = ["Percentage"] + [f"{pct}%" for _, pct in coverage]
        col_widths = [3.5 * cm] + [
            (16.0 - 3.5) / len(coverage) * cm for _ in coverage
        ]
        table = Table([header_row, pct_row], colWidths=col_widths)
        table.setStyle(
            TableStyle([
                ("BOX", (0, 0), (-1, -1), 0.75, colors.black),
                ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.black),
                ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9.5),
                ("ALIGN", (1, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ])
        )
        return [Spacer(1, 8), KeepTogether([table]), Spacer(1, 8)]

    @staticmethod
    def _bloom_short_code(bloom_level: str) -> str:
        """Map a full Bloom's level name to its short "L1"-"L6" code."""
        order = [
            "Remember", "Understand", "Apply", "Analyze", "Evaluate", "Create",
        ]
        try:
            return f"L{order.index(bloom_level) + 1}"
        except ValueError:
            return bloom_level[:2] if bloom_level else "—"

    # ── Page callback ─────────────────────────────────────────────────────────

    def _make_page_callback(self, pa: ParsedAssessment):
        """Return an onPage callback that draws the page header and footer.

        Args:
            pa: ParsedAssessment supplying header/footer text.

        Returns:
            Callable[[canvas, doc], None] for SimpleDocTemplate.
        """
        navy = self._NAVY
        grey = self._GREY_TEXT
        title = pa.title

        def _draw_page(canvas, doc):
            canvas.saveState()
            w, h = A4

            # ── Top rule ────────────────────────────────────────────────────
            canvas.setStrokeColor(navy)
            canvas.setLineWidth(1.5)
            canvas.line(2.5 * cm, h - 1.5 * cm, w - 2.5 * cm, h - 1.5 * cm)

            # ── Footer rule + text ───────────────────────────────────────────
            canvas.setLineWidth(0.5)
            canvas.line(2.5 * cm, 1.8 * cm, w - 2.5 * cm, 1.8 * cm)

            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(grey)

            footer_left = title
            footer_right = f"Page {doc.page}"
            canvas.drawString(2.5 * cm, 1.3 * cm, footer_left)
            canvas.drawRightString(w - 2.5 * cm, 1.3 * cm, footer_right)

            canvas.restoreState()

        return _draw_page

    # ── Utilities ─────────────────────────────────────────────────────────────

    @staticmethod
    def _esc(text: str) -> str:
        """Escape HTML special characters to prevent ReportLab Paragraph errors.

        Replaces ``&``, ``<``, and ``>`` with their HTML entity equivalents
        so that raw LLM output never crashes the paragraph parser.

        Args:
            text: Raw text string.

        Returns:
            str: HTML-safe string for ReportLab Paragraph().
        """
        if not text:
            return ""
        text = text.replace("&", "&amp;")
        text = text.replace("<", "&lt;")
        text = text.replace(">", "&gt;")
        return text

    @classmethod
    def _accreditation_html(cls, line: str) -> str:
        """Wrap key accreditation terms in red <font> tags for ReportLab.

        Splits *line* on the shared highlight-word pattern (see
        ``_ACCREDITATION_HIGHLIGHT_RE`` / ``_ACCREDITATION_HIGHLIGHT_WORDS``
        in the module-level constants) and marks each matched word in red,
        matching the Word exporter's run-colouring and the official
        paper's styling.

        Args:
            line: One raw accreditation line (e.g. "Approved by AICTE").

        Returns:
            str: HTML-safe, ReportLab-renderable markup with highlighted
            words wrapped in ``<font color="#CC0000">...</font>``.
        """
        parts = _ACCREDITATION_HIGHLIGHT_RE.split(line)
        out = []
        for part in parts:
            if not part:
                continue
            escaped = cls._esc(part)
            if part in _ACCREDITATION_HIGHLIGHT_WORDS:
                out.append(
                    f'<font color="#{_ACCREDITATION_HIGHLIGHT_COLOR}">{escaped}</font>'
                )
            else:
                out.append(escaped)
        return "".join(out)


# ===========================================================================
# Convenience factory
# ===========================================================================


class DownloadEngine:
    """Facade that wraps all three exporters behind a single interface.

    Instantiate once and call the appropriate ``export_*`` method. All
    exceptions from the underlying exporters are caught, logged, and
    re-raised so callers receive a descriptive error message.

    Example::

        engine = DownloadEngine()
        md_text = engine.export_markdown(assessment)
        docx_bytes = engine.export_word(assessment)
        pdf_bytes = engine.export_pdf(assessment)
    """

    def __init__(self) -> None:
        self._md = MarkdownExporter()
        self._word = WordExporter()
        self._pdf = PDFExporter()

    def export_markdown(self, assessment: Assessment) -> str:
        """Export *assessment* to Markdown.

        Args:
            assessment: The Assessment to export.

        Returns:
            str: Markdown document.

        Raises:
            RuntimeError: If export fails.
        """
        try:
            return self._md.export(assessment)
        except Exception as exc:
            logger.exception("MarkdownExporter failed: %s", exc)
            raise RuntimeError(f"Markdown export failed: {exc}") from exc

    def export_word(self, assessment: Assessment) -> bytes:
        """Export *assessment* to a DOCX byte stream.

        Args:
            assessment: The Assessment to export.

        Returns:
            bytes: Raw .docx bytes.

        Raises:
            RuntimeError: If export fails.
        """
        try:
            return self._word.export(assessment)
        except Exception as exc:
            logger.exception("WordExporter failed: %s", exc)
            raise RuntimeError(f"Word export failed: {exc}") from exc

    def export_pdf(self, assessment: Assessment) -> bytes:
        """Export *assessment* to a PDF byte stream.

        Args:
            assessment: The Assessment to export.

        Returns:
            bytes: Raw PDF bytes.

        Raises:
            RuntimeError: If export fails.
        """
        try:
            return self._pdf.export(assessment)
        except Exception as exc:
            logger.exception("PDFExporter failed: %s", exc)
            raise RuntimeError(f"PDF export failed: {exc}") from exc
